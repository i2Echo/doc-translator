"""Single-DOCX regression runner for the independent DOCX translation path.

A repeatable command that translates one DOCX through the same module the
worker uses (``docx_translator.translate_docx``) -- *not* the HTTP API -- and
emits, into a run directory:

* ``output.docx``      -- the translated DOCX
* ``format-report.json`` -- per-paragraph format and structure report

By default the runner uses a deterministic ``MockTranslator`` (no model API
required) that prefixes translated text with a target-language marker so
assertions are reproducible. Pass ``--model-*`` flags to use a real
OpenAI-compatible model.

The runner also embeds a small ``--self-test`` mode that generates sample
DOCX files covering mixed formatting, hyperlinks, drawings, controls, fields,
tables and merged cells, then verifies preservation end to end.

Usage::

    py -3.12 -m tests.regression.run_single_docx \
        --input tests/regression/inputs/sample.docx \
        --output-dir tests/regression/runs/docx-sample

    py -3.12 -m tests.regression.run_single_docx --self-test \
        --output-dir tests/regression/runs/docx-self-test
"""

from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass
from pathlib import Path

# Make the api package importable when run from the repo root.
_REPO_ROOT = Path(__file__).resolve().parents[2]
_API_ROOT = _REPO_ROOT / "apps" / "api"
if str(_API_ROOT) not in sys.path:
    sys.path.insert(0, str(_API_ROOT))

from docx import Document  # noqa: E402
from docx.document import Document as DocumentType  # noqa: E402
from docx.opc.constants import RELATIONSHIP_TYPE  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402
from PIL import Image  # noqa: E402

from doc_translator.docx_translator import _collect_targets, translate_docx  # noqa: E402


class MockTranslator:
    """Returns a deterministic transformation of the input text.

    The prefix ``[ZH]`` is prepended so translated content is easy to verify
    in assertions. Whitespace-only input is returned unchanged to mirror the
    real translator's behaviour.
    """

    def translate_text(
        self,
        text: str,
        *,
        source_language: str,
        target_language: str,
        preserve_line_breaks: bool = True,
        extra_system_instruction: str = "",
    ) -> str:
        if not text.strip():
            return text
        return f"[{target_language[:2].upper()}]{text}"


def _translate_segments(
    translator,
    segments: list[str],
    *,
    source_language: str,
    target_language: str,
    preserve_line_breaks: bool,
    on_progress,
    cancel_check,
) -> list[str]:
    translated: list[str] = []
    total = len(segments)
    for index, segment in enumerate(segments, start=1):
        cancel_check()
        translated.append(
            translator.translate_text(
                segment,
                source_language=source_language,
                target_language=target_language,
                preserve_line_breaks=preserve_line_breaks,
            )
        )
        on_progress(index, total)
    return translated


# ---------------------------------------------------------------------------
# Format report
# ---------------------------------------------------------------------------


@dataclass
class RunFormat:
    text: str
    bold: bool | None
    italic: bool | None
    underline: bool | None
    strike: bool | None
    name: str | None
    size: float | None
    color: str | None


def _run_format(run) -> RunFormat:
    color = run.font.color.rgb
    return RunFormat(
        text=run.text,
        bold=run.font.bold,
        italic=run.font.italic,
        underline=run.font.underline,
        strike=run.font.strike,
        name=run.font.name,
        size=run.font.size,
        color=str(color) if color is not None else None,
    )


def _build_format_report(input_path: Path, output_path: Path) -> dict:
    """Compare run-level formatting between input and output documents."""
    in_doc = Document(input_path)
    out_doc = Document(output_path)
    in_targets = _collect_targets(in_doc)
    out_targets = _collect_targets(out_doc)

    paragraphs: list[dict] = []
    all_preserved = True
    all_structures_preserved = True
    for index, (in_para, out_para) in enumerate(zip(in_targets, out_targets, strict=True), start=1):
        in_runs = [_run_format(r) for r in in_para.runs]
        out_runs = [_run_format(r) for r in out_para.runs]
        # Compare format signatures (excluding text content) pairwise.
        format_preserved = len(in_runs) == len(out_runs) and all(
            _format_equal(a, b) for a, b in zip(in_runs, out_runs, strict=True)
        )
        if not format_preserved:
            all_preserved = False
        input_structure = _structure_counts(in_para)
        output_structure = _structure_counts(out_para)
        structure_preserved = input_structure == output_structure
        if not structure_preserved:
            all_structures_preserved = False
        paragraphs.append(
            {
                "index": index,
                "format_preserved": format_preserved,
                "structure_preserved": structure_preserved,
                "input_text": _visible_text(in_para),
                "output_text": _visible_text(out_para),
                "input_structure": input_structure,
                "output_structure": output_structure,
                "input_runs": [r.__dict__ for r in in_runs],
                "output_runs": [r.__dict__ for r in out_runs],
            }
        )

    return {
        "input_path": str(input_path),
        "output_path": str(output_path),
        "input_paragraph_count": len(in_targets),
        "output_paragraph_count": len(out_targets),
        "all_formats_preserved": all_preserved,
        "all_structures_preserved": all_structures_preserved,
        "paragraphs": paragraphs,
    }


def _format_equal(a: RunFormat, b: RunFormat) -> bool:
    return (
        a.bold == b.bold
        and a.italic == b.italic
        and a.underline == b.underline
        and a.strike == b.strike
        and a.name == b.name
        and a.size == b.size
        and a.color == b.color
    )


def _visible_text(paragraph: Paragraph) -> str:
    return "".join(element.text or "" for element in paragraph._element.iter(qn("w:t")))


def _structure_counts(paragraph: Paragraph) -> dict:
    hyperlink_targets = []
    hyperlink_formats = []
    for hyperlink in paragraph._element.iter(qn("w:hyperlink")):
        relationship_id = hyperlink.get(qn("r:id"))
        hyperlink_targets.append(paragraph.part.rels[relationship_id].target_ref if relationship_id else None)
        hyperlink_formats.extend(
            properties.xml
            for properties in hyperlink.iter(qn("w:rPr"))
        )
    return {
        "hyperlinks": len(list(paragraph._element.iter(qn("w:hyperlink")))),
        "hyperlink_targets": hyperlink_targets,
        "hyperlink_formats": hyperlink_formats,
        "drawings": len(list(paragraph._element.iter(qn("w:drawing")))),
        "tabs": len(list(paragraph._element.iter(qn("w:tab")))),
        "breaks": len(list(paragraph._element.iter(qn("w:br")))),
        "field_chars": len(list(paragraph._element.iter(qn("w:fldChar")))),
    }


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_field(paragraph: Paragraph, instruction: str, result: str) -> None:
    begin_run = paragraph.add_run()._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)

    instruction_run = paragraph.add_run()._r
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_element.text = instruction
    instruction_run.append(instruction_element)

    separate_run = paragraph.add_run()._r
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph.add_run(result)

    end_run = paragraph.add_run()._r
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)


# ---------------------------------------------------------------------------
# Self-test: generate sample DOCX files covering edge cases
# ---------------------------------------------------------------------------


def _generate_self_test_docx(path: Path) -> None:
    """Create a DOCX with the formatting scenarios the regression suite cares about."""
    doc = Document()

    # 1. Single run plain text.
    doc.add_paragraph("Hello world single run.")

    # 2. Multiple runs, uniform formatting -> should collapse to first run.
    p2 = doc.add_paragraph()
    p2.add_run("Hello ")
    p2.add_run("world ")
    p2.add_run("uniform runs.")

    # 3. Mixed formatting: bold + normal -> runs must keep their bold flag.
    p3 = doc.add_paragraph()
    bold_run = p3.add_run("Bold part here")
    bold_run.font.bold = True
    normal_run = p3.add_run(" and normal part here.")
    normal_run.font.bold = False

    # 4. Table with translatable cell paragraphs.
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Cell A text")
    table.rows[0].cells[1].paragraphs[0].add_run("Cell B text")

    # 5. Paragraph after table.
    doc.add_paragraph("After table paragraph.")

    # 6. A hyperlink between ordinary runs must stay in document order.
    hyperlink_paragraph = doc.add_paragraph()
    hyperlink_paragraph.add_run("Before ")
    _add_hyperlink(hyperlink_paragraph, "linked text", "https://example.com")
    hyperlink_paragraph.add_run(" after")

    # 7. Replacing surrounding text must not delete an inline drawing.
    image_path = path.with_suffix(".png")
    Image.new("RGB", (2, 2), "white").save(image_path)
    drawing_paragraph = doc.add_paragraph()
    drawing_paragraph.add_run("Before image ")
    drawing_paragraph.add_run().add_picture(str(image_path))
    drawing_paragraph.add_run(" after image")

    # 8. Tabs and manual line breaks remain structural elements.
    controls_paragraph = doc.add_paragraph()
    controls_run = controls_paragraph.add_run("Before tab")
    controls_run.add_tab()
    controls_run.add_text("after tab")
    controls_run.add_break()
    controls_run.add_text("after break")

    # 9. Field instructions remain intact while their displayed result is translated.
    field_paragraph = doc.add_paragraph("Document date: ")
    _add_field(field_paragraph, " DATE ", "July 13, 2026")

    # 10. A merged cell is exposed more than once by python-docx but translated once.
    merged_table = doc.add_table(rows=1, cols=2)
    merged_cell = merged_table.cell(0, 0).merge(merged_table.cell(0, 1))
    merged_cell.paragraphs[0].add_run("Merged cell text")

    # 11. Empty paragraph -> should be skipped by _collect_targets.
    doc.add_paragraph("")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    image_path.unlink()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def _run_translation(input_path: Path, output_path: Path, args: argparse.Namespace) -> list[str]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    events: list[str] = []

    if args.use_real_model:
        from doc_translator.settings_service import RuntimeSettings
        from doc_translator.translation import OpenAICompatibleTranslator, translate_segments

        runtime = RuntimeSettings(
            storage_mode="local",
            local_storage_path=str(output_path.parent),
            file_retention_days=7,
            model_base_url=args.model_base_url,
            model_api_key=args.model_api_key,
            model_name=args.model_name,
            model_timeout_seconds=args.model_timeout_seconds,
            ocr_enabled=False,
            ocr_language_hint="auto",
            max_upload_mb=100,
            max_concurrent_jobs=1,
        )
        translator = OpenAICompatibleTranslator(runtime)
        segment_translator = translate_segments
    else:
        translator = MockTranslator()
        segment_translator = _translate_segments

    translate_docx(
        str(input_path),
        output_path,
        translator=translator,
        translate_segments=segment_translator,
        source_language=args.source_language,
        target_language=args.target_language,
        on_progress=lambda index, total: events.append(f"progress:{index}/{total}"),
        cancel_check=lambda: events.append("cancel-check"),
        on_rebuilding=lambda: events.append("rebuilding"),
    )
    return events


def _assert_no_partial_output(output_path: Path) -> None:
    if output_path.exists():
        raise AssertionError(f"Unexpected final output: {output_path}")
    temp_files = list(output_path.parent.glob(f".{output_path.name}.*.tmp"))
    if temp_files:
        raise AssertionError(f"Temporary DOCX files were not cleaned: {temp_files}")


def _verify_failure_paths(input_path: Path, output_dir: Path) -> None:
    callbacks = {
        "translator": MockTranslator(),
        "source_language": "English",
        "target_language": "Chinese",
        "on_progress": lambda index, total: None,
        "on_rebuilding": lambda: None,
    }

    mismatch_output = output_dir / "mismatch.docx"
    try:
        translate_docx(
            str(input_path),
            mismatch_output,
            translate_segments=lambda translator, segments, **kwargs: [],
            cancel_check=lambda: None,
            **callbacks,
        )
    except RuntimeError as exc:
        if "results" not in str(exc) or "text spans" not in str(exc):
            raise
    else:
        raise AssertionError("Mismatched translation count did not fail")
    _assert_no_partial_output(mismatch_output)

    model_failure_output = output_dir / "model-failure.docx"

    def fail_translation(translator, segments, **kwargs):
        raise RuntimeError("simulated model failure")

    try:
        translate_docx(
            str(input_path),
            model_failure_output,
            translate_segments=fail_translation,
            cancel_check=lambda: None,
            **callbacks,
        )
    except RuntimeError as exc:
        if str(exc) != "simulated model failure":
            raise
    else:
        raise AssertionError("Model failure did not propagate")
    _assert_no_partial_output(model_failure_output)

    cancelled_output = output_dir / "cancelled.docx"
    cancel_checks = 0

    def cancel_after_validation() -> None:
        nonlocal cancel_checks
        cancel_checks += 1
        if cancel_checks == 2:
            raise RuntimeError("cancelled after validation")

    try:
        translate_docx(
            str(input_path),
            cancelled_output,
            translate_segments=lambda translator, segments, **kwargs: list(segments),
            cancel_check=cancel_after_validation,
            **callbacks,
        )
    except RuntimeError as exc:
        if str(exc) != "cancelled after validation":
            raise
    else:
        raise AssertionError("Cancellation before publication did not fail")
    _assert_no_partial_output(cancelled_output)

    save_failure_output = output_dir / "save-failure.docx"
    original_save = DocumentType.save

    def fail_after_save(document, target) -> None:
        original_save(document, target)
        raise OSError("simulated save failure")

    DocumentType.save = fail_after_save
    try:
        try:
            translate_docx(
                str(input_path),
                save_failure_output,
                translate_segments=lambda translator, segments, **kwargs: list(segments),
                cancel_check=lambda: None,
                **callbacks,
            )
        except OSError as exc:
            if str(exc) != "simulated save failure":
                raise
        else:
            raise AssertionError("Save failure did not propagate")
    finally:
        DocumentType.save = original_save
    _assert_no_partial_output(save_failure_output)

    empty_input = output_dir / "empty-input.docx"
    empty_output = output_dir / "empty-output.docx"
    Document().save(empty_input)
    translator_called = False

    def reject_empty_translation(translator, segments, **kwargs):
        nonlocal translator_called
        translator_called = True
        raise AssertionError("Translator was called for an empty DOCX")

    translate_docx(
        str(empty_input),
        empty_output,
        translate_segments=reject_empty_translation,
        cancel_check=lambda: None,
        **callbacks,
    )
    if translator_called:
        raise AssertionError("Translator was called for an empty DOCX")
    Document(empty_output)
    if list(empty_output.parent.glob(f".{empty_output.name}.*.tmp")):
        raise AssertionError("Empty DOCX translation left a temporary file")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run a single DOCX through the independent DOCX translation path.")
    parser.add_argument("--input", type=Path, default=None, help="Input DOCX path. Required unless --self-test.")
    parser.add_argument("--output-dir", required=True, type=Path, help="Run output directory.")
    parser.add_argument("--source-language", default="English")
    parser.add_argument("--target-language", default="Chinese")
    parser.add_argument("--self-test", action="store_true", help="Generate sample DOCX files and verify format preservation.")
    parser.add_argument("--use-real-model", action="store_true", help="Use a real OpenAI-compatible model instead of the mock translator.")
    parser.add_argument("--model-base-url", default="https://api.openai.com/v1")
    parser.add_argument("--model-api-key", default="")
    parser.add_argument("--model-name", default="gpt-4.1-mini")
    parser.add_argument("--model-timeout-seconds", type=int, default=120)
    args = parser.parse_args(argv)

    output_dir: Path = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    if args.self_test:
        input_path = output_dir / "self-test-input.docx"
        _generate_self_test_docx(input_path)
        print(f"Generated self-test input: {input_path}")
    else:
        if args.input is None:
            parser.error("--input is required unless --self-test is given")
        input_path = args.input

    output_path = output_dir / "output.docx"
    events = _run_translation(input_path, output_path, args)
    print(f"Translated DOCX written: {output_path}")

    report = _build_format_report(input_path, output_path)
    report_path = output_dir / "format-report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Format report written: {report_path}")

    print()
    print(f"Input paragraphs : {report['input_paragraph_count']}")
    print(f"Output paragraphs: {report['output_paragraph_count']}")
    print(f"Formats preserved: {report['all_formats_preserved']}")
    print(f"Structures preserved: {report['all_structures_preserved']}")

    if args.self_test:
        # In self-test mode, preservation must hold and counts must match.
        if not report["all_formats_preserved"]:
            print("SELF-TEST FAILED: run-level formatting was not preserved")
            return 1
        if report["input_paragraph_count"] != report["output_paragraph_count"]:
            print("SELF-TEST FAILED: paragraph count mismatch")
            return 1
        if not report["all_structures_preserved"]:
            print("SELF-TEST FAILED: inline structure was not preserved")
            return 1
        if not any(event.startswith("progress:1/") for event in events) or events[-2:] != [
            "rebuilding",
            "cancel-check",
        ]:
            print("SELF-TEST FAILED: progress, cancellation, or rebuilding callback order changed")
            return 1
        if not args.use_real_model:
            hyperlink_output = next(
                paragraph["output_text"]
                for paragraph in report["paragraphs"]
                if paragraph["input_text"] == "Before linked text after"
            )
            if not (
                hyperlink_output.index("Before")
                < hyperlink_output.index("linked text")
                < hyperlink_output.index("after")
            ):
                print("SELF-TEST FAILED: hyperlink text order changed")
                return 1
            merged_output = next(
                paragraph["output_text"]
                for paragraph in report["paragraphs"]
                if paragraph["input_text"] == "Merged cell text"
            )
            if merged_output.count("[CH]") != 1:
                print("SELF-TEST FAILED: merged cell was translated more than once")
                return 1
        try:
            _verify_failure_paths(input_path, output_dir)
        except AssertionError as exc:
            print(f"SELF-TEST FAILED: {exc}")
            return 1
        print("SELF-TEST PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main())
