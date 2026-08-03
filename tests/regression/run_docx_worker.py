"""Exercise DOCX jobs through run_translation_job without external services."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from uuid import uuid4

_REPO_ROOT = Path(__file__).resolve().parents[2]
_API_ROOT = _REPO_ROOT / "apps" / "api"
if str(_API_ROOT) not in sys.path:
    sys.path.insert(0, str(_API_ROOT))

from docx import Document  # noqa: E402
from sqlalchemy import create_engine, select  # noqa: E402
from sqlalchemy.orm import sessionmaker  # noqa: E402

import doc_translator.translation as translation  # noqa: E402
from doc_translator.db import Base  # noqa: E402
from doc_translator.model_api import ModelApiFormat  # noqa: E402
from doc_translator.models import JobEvent, JobFile, JobFileKind, JobStatus, TranslationJob, User, UserRole  # noqa: E402
from doc_translator.preview import preview_sidecar_path  # noqa: E402
from doc_translator.settings_service import RuntimeSettings  # noqa: E402


class MockTranslator:
    def close(self) -> None:
        pass

    def translate_text(
        self,
        text: str,
        *,
        source_language: str,
        target_language: str,
        preserve_line_breaks: bool = True,
        extra_system_instruction: str = "",
    ) -> str:
        return f"[{target_language[:2].upper()}]{text}"


def _create_job(session_factory, storage_root: Path, user_id: str, label: str) -> str:
    input_path = storage_root / "uploads" / f"{label}.docx"
    document = Document()
    document.add_paragraph(f"{label} first paragraph")
    document.add_paragraph(f"{label} second paragraph")
    document.save(input_path)

    with session_factory() as session:
        input_file = JobFile(
            original_name=f"{label}.docx",
            stored_name=input_path.name,
            storage_path=str(input_path),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=input_path.stat().st_size,
            checksum="0" * 64,
            kind=JobFileKind.INPUT,
            created_by=user_id,
        )
        job = TranslationJob(
            created_by=user_id,
            status=JobStatus.QUEUED,
            progress=0,
            source_language="English",
            target_language="Chinese",
            input_file=input_file,
            model_base_url_snapshot="https://example.com/v1",
            model_name_snapshot="mock",
            model_api_format_snapshot=ModelApiFormat.CHAT_COMPLETIONS.value,
        )
        session.add(job)
        session.commit()
        return job.id


def _load_job(session_factory, job_id: str) -> TranslationJob:
    with session_factory() as session:
        job = session.get(TranslationJob, job_id)
        if job is None:
            raise AssertionError(f"Job disappeared: {job_id}")
        session.expunge(job)
        return job


def _output_paths(session_factory, job_id: str) -> list[Path]:
    with session_factory() as session:
        return [
            Path(path)
            for path in session.scalars(
                select(JobFile.storage_path).where(
                    JobFile.kind == JobFileKind.OUTPUT,
                    TranslationJob.id == job_id,
                    TranslationJob.output_file_id == JobFile.id,
                )
            )
        ]


def _assert_no_temporary_files(storage_root: Path) -> None:
    temporary_files = list((storage_root / "results").glob(".*.tmp"))
    if temporary_files:
        raise AssertionError(f"Temporary result files remain: {temporary_files}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run DOCX worker-path regression scenarios.")
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args(argv)

    output_dir: Path = args.output_dir
    storage_root = output_dir / "files"
    (storage_root / "uploads").mkdir(parents=True, exist_ok=True)
    (storage_root / "results").mkdir(parents=True, exist_ok=True)
    database_path = output_dir / "worker-regression.sqlite3"
    database_path.unlink(missing_ok=True)

    engine = create_engine(f"sqlite+pysqlite:///{database_path}")
    session_factory = sessionmaker(bind=engine, autoflush=False, autocommit=False, expire_on_commit=False)
    Base.metadata.create_all(engine)

    runtime = RuntimeSettings(
        storage_mode="local",
        local_storage_path=str(storage_root),
        file_retention_days=7,
        model_api_format=ModelApiFormat.CHAT_COMPLETIONS,
        model_base_url="https://example.com/v1",
        model_api_key="",
        model_name="mock",
        model_timeout_seconds=30,
        ocr_enabled=False,
        ocr_language_hint="auto",
        max_upload_mb=10,
        max_concurrent_jobs=1,
    )

    user_id = str(uuid4())
    with session_factory() as session:
        session.add(
            User(
                id=user_id,
                email="docx-worker@example.com",
                full_name="DOCX Worker Regression",
                password_hash="unused",
                role=UserRole.ADMIN,
            )
        )
        session.commit()

    original_session_local = translation.SessionLocal
    original_runtime_loader = translation.get_runtime_settings
    original_translator = translation.ModelTranslator
    original_preview_loader = translation.load_or_create_preview
    original_checksum = translation.file_checksum
    translation.SessionLocal = session_factory
    translation.get_runtime_settings = lambda session: runtime
    translation.ModelTranslator = lambda runtime: MockTranslator()

    try:
        success_job_id = _create_job(session_factory, storage_root, user_id, "success")
        translation.run_translation_job(success_job_id)
        success_job = _load_job(session_factory, success_job_id)
        success_outputs = _output_paths(session_factory, success_job_id)
        if success_job.status != JobStatus.COMPLETED or len(success_outputs) != 1:
            raise AssertionError("Successful DOCX job did not complete with one output")
        output_document = Document(success_outputs[0])
        if not all(paragraph.text.startswith("[CH]") for paragraph in output_document.paragraphs):
            raise AssertionError("Worker output did not contain translated text")
        if not preview_sidecar_path(str(success_outputs[0])).exists():
            raise AssertionError("Successful DOCX job did not create its preview")

        preview_job_id = _create_job(session_factory, storage_root, user_id, "preview-failure")

        def fail_preview(job, *, force=False):
            raise RuntimeError("simulated preview failure")

        translation.load_or_create_preview = fail_preview
        translation.run_translation_job(preview_job_id)
        preview_job = _load_job(session_factory, preview_job_id)
        if preview_job.status != JobStatus.COMPLETED or len(_output_paths(session_factory, preview_job_id)) != 1:
            raise AssertionError("DOCX preview failure did not degrade to a completed job")
        with session_factory() as session:
            preview_events = session.scalars(
                select(JobEvent.message).where(JobEvent.job_id == preview_job_id)
            ).all()
        if "Preview could not be prepared" not in preview_events:
            raise AssertionError("DOCX preview failure was not recorded")
        translation.load_or_create_preview = original_preview_loader

        registration_job_id = _create_job(session_factory, storage_root, user_id, "registration-failure")
        existing_results = set((storage_root / "results").iterdir())

        def fail_checksum(path: Path) -> str:
            raise RuntimeError("simulated output registration failure")

        translation.file_checksum = fail_checksum
        translation.run_translation_job(registration_job_id)
        registration_job = _load_job(session_factory, registration_job_id)
        if registration_job.status != JobStatus.FAILED or registration_job.output_file_id is not None:
            raise AssertionError("Output registration failure did not fail cleanly")
        if set((storage_root / "results").iterdir()) != existing_results:
            raise AssertionError("Output registration failure left an orphan result")
        translation.file_checksum = original_checksum

        cancellation_job_id = _create_job(session_factory, storage_root, user_id, "cancellation")

        class CancellingTranslator(MockTranslator):
            cancelled = False

            def translate_text(self, text: str, **kwargs) -> str:
                translated_text = super().translate_text(text, **kwargs)
                if not self.cancelled:
                    self.cancelled = True
                    with session_factory() as cancel_session:
                        job = cancel_session.get(TranslationJob, cancellation_job_id)
                        job.cancel_requested = True
                        cancel_session.commit()
                return translated_text

        translation.ModelTranslator = lambda runtime: CancellingTranslator()
        existing_results = set((storage_root / "results").iterdir())
        translation.run_translation_job(cancellation_job_id)
        cancellation_job = _load_job(session_factory, cancellation_job_id)
        if cancellation_job.status != JobStatus.CANCELLED or cancellation_job.output_file_id is not None:
            raise AssertionError("Cancelled DOCX job did not stop cleanly")
        if set((storage_root / "results").iterdir()) != existing_results:
            raise AssertionError("Cancelled DOCX job left a result")

        _assert_no_temporary_files(storage_root)
    finally:
        translation.SessionLocal = original_session_local
        translation.get_runtime_settings = original_runtime_loader
        translation.ModelTranslator = original_translator
        translation.load_or_create_preview = original_preview_loader
        translation.file_checksum = original_checksum
        engine.dispose()

    print("DOCX WORKER REGRESSION PASSED")
    print("Scenarios: success, preview degradation, registration failure cleanup, cancellation cleanup")
    return 0


if __name__ == "__main__":
    sys.exit(main())
