from __future__ import annotations

import logging
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Callable

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from doc_translator.translation import OpenAICompatibleTranslator

logger = logging.getLogger(__name__)

_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


@dataclass
class _TextSpan:
    elements: list

    @property
    def text(self) -> str:
        return "".join(element.text or "" for element in self.elements)

    @property
    def translatable_text(self) -> str:
        return self.text.strip()

    def replace(self, text: str) -> None:
        original = self.text
        leading_length = len(original) - len(original.lstrip())
        trailing_length = len(original) - len(original.rstrip())
        leading = original[:leading_length]
        trailing = original[len(original) - trailing_length:] if trailing_length else ""
        first, *rest = self.elements
        _set_text(first, f"{leading}{text}{trailing}")
        for element in rest:
            _set_text(element, "")


def _set_text(element, text: str) -> None:
    element.text = text
    if text[:1].isspace() or text[-1:].isspace():
        element.set(_XML_SPACE, "preserve")
    else:
        element.attrib.pop(_XML_SPACE, None)


def _run_format_key(run_element) -> str:
    properties = run_element.find(qn("w:rPr"))
    return properties.xml if properties is not None else ""


def _collect_text_spans(paragraph: Paragraph) -> list[_TextSpan]:
    spans: list[_TextSpan] = []
    active_span: _TextSpan | None = None
    active_key: tuple[int | None, str] | None = None

    def reset() -> None:
        nonlocal active_span, active_key
        active_span = None
        active_key = None

    def visit_run(run_element, container_id: int | None) -> None:
        nonlocal active_span, active_key
        key = (container_id, _run_format_key(run_element))
        for child in run_element.iterchildren():
            if child.tag == qn("w:rPr"):
                continue
            if child.tag != qn("w:t"):
                reset()
                continue
            if not child.text:
                continue
            if active_span is None or active_key != key:
                active_span = _TextSpan([])
                spans.append(active_span)
                active_key = key
            active_span.elements.append(child)

    for child in paragraph._element.iterchildren():
        if child.tag == qn("w:r"):
            visit_run(child, None)
            continue
        if child.tag == qn("w:hyperlink"):
            reset()
            container_id = id(child)
            for hyperlink_child in child.iterchildren():
                if hyperlink_child.tag == qn("w:r"):
                    visit_run(hyperlink_child, container_id)
                else:
                    reset()
            reset()
            continue
        reset()

    return spans


def _collect_targets(document: DocumentType) -> list[Paragraph]:
    targets: list[Paragraph] = []
    seen: set[int] = set()

    def add(paragraphs) -> None:
        for paragraph in paragraphs:
            identity = id(paragraph._element)
            if identity in seen:
                continue
            spans = _collect_text_spans(paragraph)
            if not any(span.translatable_text for span in spans):
                continue
            seen.add(identity)
            targets.append(paragraph)

    add(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                add(cell.paragraphs)
    return targets


def translate_docx(
    input_path: str,
    output_path: Path,
    *,
    translator: OpenAICompatibleTranslator,
    translate_segments: Callable[..., list[str]],
    source_language: str,
    target_language: str,
    on_progress: Callable[[int, int], None],
    cancel_check: Callable[[], None],
    on_rebuilding: Callable[[], None],
) -> int | None:
    started_at = time.perf_counter()
    document = Document(input_path)
    targets = _collect_targets(document)
    spans = [
        span
        for paragraph in targets
        for span in _collect_text_spans(paragraph)
        if span.translatable_text
    ]

    if spans:
        translated = translate_segments(
            translator,
            [span.translatable_text for span in spans],
            source_language=source_language,
            target_language=target_language,
            preserve_line_breaks=True,
            on_progress=on_progress,
            cancel_check=cancel_check,
        )
    else:
        translated = []
        logger.info("DOCX contains no translatable text", extra={"input": input_path})

    if len(translated) != len(spans):
        raise RuntimeError(
            f"DOCX translation returned {len(translated)} results for {len(spans)} text spans"
        )

    for span, translated_text in zip(spans, translated, strict=True):
        span.replace(translated_text)

    cancel_check()
    on_rebuilding()
    temp_path: Path | None = None
    save_started_at = time.perf_counter()
    try:
        with tempfile.NamedTemporaryFile(
            dir=output_path.parent,
            prefix=f".{output_path.name}.",
            suffix=".tmp",
            delete=False,
        ) as temp_file:
            temp_path = Path(temp_file.name)
        document.save(temp_path)
        Document(temp_path)
        cancel_check()
        temp_path.replace(output_path)
    finally:
        if temp_path is not None:
            temp_path.unlink(missing_ok=True)

    logger.info(
        "Saved translated DOCX",
        extra={
            "output": str(output_path),
            "paragraphs": len(targets),
            "spans": len(spans),
            "save_seconds": round(time.perf_counter() - save_started_at, 3),
            "total_seconds": round(time.perf_counter() - started_at, 3),
        },
    )
    return None
