from __future__ import annotations

import json
import re
import unicodedata
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime, timezone
from difflib import SequenceMatcher
from functools import lru_cache
from pathlib import Path
from tempfile import NamedTemporaryFile

import fitz
from docx import Document

from doc_translator.interceptors.alignment_sniffer import sniff_alignment
from doc_translator.interceptors.grouper import dominant_value
from doc_translator.interceptors.weight_detector import detect_font_style
from doc_translator.render_guards.cascade_scaler import minimum_sibling_font_size
from doc_translator.render_guards.font_router import font_route_for_language
from doc_translator.render_guards.line_wrapper import apply_thai_word_wrap_shield
from doc_translator.models import TranslationJob
from doc_translator.storage import file_checksum


DOCX_PREVIEW_PARAGRAPH_LIMIT = 8
DOCX_PREVIEW_CHAR_LIMIT = 2200
PREVIEW_SCHEMA_VERSION = 9
PDF_PREVIEW_TEXT_GRANULARITY = "visual-paragraph"
PDF_CJK_SERIF_FONTFILES = (
    "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
    "/usr/share/fonts/noto/NotoSerifCJK-Regular.ttc",
    "/usr/share/fonts/google-noto-cjk/NotoSerifCJK-Regular.ttc",
    r"C:\Windows\Fonts\simsun.ttc",
    r"C:\Windows\Fonts\msyh.ttc",
)
PDF_CJK_SANS_FONTFILES = (
    "/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.ttf",
    "/usr/share/fonts/noto/NotoSansCJKsc-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    r"C:\Windows\Fonts\NotoSansSC-VF.ttf",
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\msyhbd.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
)
PDF_LATIN_FONTFILES = (
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    "/usr/share/fonts/noto/NotoSans-Regular.ttf",
    "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    r"C:\Windows\Fonts\arial.ttf",
)
PDF_THAI_FONTFILES = (
    "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
    "/usr/share/fonts/noto/NotoSansThai-Regular.ttf",
    r"C:\Windows\Fonts\NotoSansThai-Regular.ttf",
    r"C:\Windows\Fonts\LeelawUI.ttf",
)
PDF_CJK_JP_FONTFILES = (
    "/usr/share/fonts/truetype/noto/NotoSansCJKjp-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJKjp-Regular.otf",
    "/usr/share/fonts/opentype/noto/NotoSansCJKjp-Regular.ttf",
    "/usr/share/fonts/noto/NotoSansCJKjp-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto/NotoSansCJK-Regular.ttc",
    r"C:\Windows\Fonts\YuGothM.ttc",
)
PDF_CJK_KR_FONTFILES = (
    "/usr/share/fonts/truetype/noto/NotoSansCJKkr-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJKkr-Regular.otf",
    "/usr/share/fonts/opentype/noto/NotoSansCJKkr-Regular.ttf",
    "/usr/share/fonts/noto/NotoSansCJKkr-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto/NotoSansCJK-Regular.ttc",
    r"C:\Windows\Fonts\malgun.ttf",
)
PDF_CJK_SERIF_CACHE_FONTFILES = (
    "SourceHanSerifCN-Regular.ttf",
    "SourceHanSerifCN-Bold.ttf",
    "NotoSerifCJK-Regular.ttc",
    "LXGWWenKaiGB-Regular.1.520.ttf",
)
PDF_CJK_SANS_CACHE_FONTFILES = (
    "NotoSansCJKsc-Regular.ttf",
    "SourceHanSansCN-Regular.ttf",
    "SourceHanSansCN-Bold.ttf",
    "NotoSansCJK-Regular.ttc",
    "LXGWWenKaiGB-Regular.1.520.ttf",
)
PDF_LATIN_CACHE_FONTFILES = (
    "NotoSans-Regular.ttf",
    "LiberationSans-Regular.ttf",
)
PDF_THAI_CACHE_FONTFILES = (
    "NotoSansThai-Regular.ttf",
)
PDF_CJK_JP_CACHE_FONTFILES = (
    "NotoSansCJKjp-Regular.ttf",
    "NotoSansCJK-Regular.ttc",
)
PDF_CJK_KR_CACHE_FONTFILES = (
    "NotoSansCJKkr-Regular.ttf",
    "NotoSansCJK-Regular.ttc",
)
PDF_CJK_SERIF_FONTNAME = "noto-serif-cjk"
PDF_CJK_SANS_FONTNAME = "noto-sans-cjk"
PDF_LATIN_FONTNAME = "noto-sans-latin"
PDF_THAI_FONTNAME = "noto-sans-thai"
PDF_CJK_JP_FONTNAME = "noto-sans-cjk-jp"
PDF_CJK_KR_FONTNAME = "noto-sans-cjk-kr"
PDF_MIN_REDRAW_FONT_SIZE = 6.0
PDF_EDITOR_MIN_FONT_SIZE = 8.0
PDF_BLOCK_MERGE_IOU_THRESHOLD = 0.3
PDF_BLOCK_MERGE_VERTICAL_DISTANCE = 5.0
PDF_BLOCK_MERGE_HORIZONTAL_GAP = 18.0
PDF_BLOCK_MERGE_MIN_HORIZONTAL_OVERLAP_RATIO = 0.2
PDF_BLOCK_GROUP_MAX_LINES = 4
PDF_BLOCK_GROUP_MAX_HEIGHT = 64.0
PDF_BLOCK_GROUP_PARAGRAPH_GAP_RATIO = 0.8
PDF_VISUAL_LINE_MIN_CENTER_TOLERANCE = 2.5
PDF_VISUAL_LINE_MAX_CENTER_TOLERANCE = 4.0
PDF_VISUAL_LINE_HEIGHT_TOLERANCE_RATIO = 0.28
PDF_MIN_VISIBLE_TEXT_SPAN_SIZE = 2.5
PDF_TABLE_MIN_ROWS = 2
PDF_TABLE_MIN_COLS = 2
PDF_TABLE_MIN_CELLS = 3
PDF_TABLE_MIN_CELL_DENSITY = 0.5
PDF_UNKNOWN_SPACING_PATTERN = re.compile(r"[\u00a0\u1680\u2000-\u200a\u202f\u205f\u3000]+")
PDF_ZERO_WIDTH_PATTERN = re.compile(r"[\u200b\u200c\u200d\ufeff]+")
PDF_CONTROL_CHAR_PATTERN = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]+")
PDF_UNKNOWN_GLYPH_BLOCK_PATTERN = re.compile(r"[\u25a0-\u25ff\u2610\u2751\u2b1a\u2b1c\ufffd]+")
PDF_MALFORMED_MULTIPLICATION_PATTERN = re.compile(r"(?<=\d)\s*[xX✕✖╳⨯*＊]\s*(?=\d)")
PDF_WRAP_TOKEN_PATTERN = re.compile(r"\s+|[A-Za-z0-9_./:%#@&+=-]+|.")
PDF_BULLET_TEXT_PATTERN = re.compile(r"^(?:\d+\s*)?[•\-\*\u2022]\s*|^\d+[\.)]\s+")
PDF_TECHNICAL_ASSIGNMENT_PATTERN = re.compile(r"^\s*([A-Z][A-Z0-9./_-]{0,10})\s*=")
PDF_FONT_SUBSET_PREFIX_PATTERN = re.compile(r"^[A-Z]{6}\+")
PDF_FONT_CAMEL_CASE_BOUNDARY_PATTERN = re.compile(r"(?<=[a-z0-9])(?=[A-Z])")
PDF_FONT_SPLIT_PATTERN = re.compile(r"[^a-z0-9]+")
PDF_FONT_IGNORED_TOKENS = frozenset({"ttf", "ttc", "otf", "psmt"})
PDF_SERIF_FONT_HINTS = ("serif", "times", "roman", "song", "ming", "mincho", "cambria", "georgia")
PDF_SANS_FONT_HINTS = ("sans", "arial", "helvetica", "gothic", "hei", "heiti", "grotesk")
PDF_MONO_FONT_HINTS = ("mono", "courier", "consola", "fixed")
PDF_SYMBOL_FONT_HINTS = ("symbol", "dingbat", "wingding")
PDF_BOLD_FONT_HINTS = ("bold", "black", "heavy", "demi", "semibold")
PDF_EMBEDDED_FONT_CHARS = frozenset("•·×−–—℃μµΩ°")
PDF_CENTER_ALIGNMENT_TOLERANCE = 5.0
PDF_ROTATION_VECTOR_TOLERANCE = 0.2


@dataclass(slots=True)
class PdfPreviewFragment:
    rect: fitz.Rect
    text: str
    font_names: list[str] = field(default_factory=list)
    font_sizes: list[float] = field(default_factory=list)
    rotations: list[int] = field(default_factory=list)

    @property
    def dominant_font(self) -> str:
        return Counter(self.font_names).most_common(1)[0][0] if self.font_names else _select_pdf_font(self.text)

    @property
    def average_font_size(self) -> float:
        return round(sum(self.font_sizes) / len(self.font_sizes), 2) if self.font_sizes else 12.0

    @property
    def is_bold(self) -> bool:
        return detect_font_style(self.font_names) == "BOLD"

    @property
    def dominant_rotation(self) -> int | None:
        rotations = [rotation for rotation in self.rotations if rotation in {90, 270}]
        return dominant_value(rotations)


@dataclass(slots=True)
class PdfPreviewWord:
    rect: fitz.Rect
    text: str
    block_index: int
    line_index: int
    word_index: int


@dataclass(frozen=True, slots=True)
class PdfPageFontResource:
    resource_name: str
    display_name: str
    tokens: tuple[str, ...]
    normalized_display_name: str
    normalized_resource_name: str
    is_bold: bool
    font_buffer: bytes | None = None


@dataclass(frozen=True, slots=True)
class PdfRenderFont:
    render_name: str
    render_file: str | None = None
    render_buffer: bytes | None = None
    metrics_name: str = "helv"
    metrics_file: str | None = None
    metrics_buffer: bytes | None = None


@dataclass(frozen=True, slots=True)
class PdfLanguageProfile:
    code: str
    font_name: str
    font_files: tuple[str, ...]
    cache_font_files: tuple[str, ...]
    min_font_size: float
    step_down: float = 0.5
    line_height_multiplier: float = 1.4
    prefer_external_font: bool = True
    aggressive_word_wrap: bool = False


PDF_LANGUAGE_ALIASES = {
    "zh": "zh",
    "zh-cn": "zh",
    "chinese": "zh",
    "simplified chinese": "zh",
    "en": "en",
    "english": "en",
    "ja": "ja",
    "japanese": "ja",
    "ko": "ko",
    "korean": "ko",
    "ms": "ms",
    "malay": "ms",
    "th": "th",
    "thai": "th",
    "vi": "vi",
    "vietnamese": "vi",
}

PDF_LANGUAGE_PROFILES = {
    "zh": PdfLanguageProfile(
        code="zh",
        font_name=PDF_CJK_SANS_FONTNAME,
        font_files=PDF_CJK_SANS_FONTFILES,
        cache_font_files=PDF_CJK_SANS_CACHE_FONTFILES,
        min_font_size=8.0,
        line_height_multiplier=1.4,
    ),
    "en": PdfLanguageProfile(
        code="en",
        font_name=PDF_LATIN_FONTNAME,
        font_files=PDF_LATIN_FONTFILES,
        cache_font_files=PDF_LATIN_CACHE_FONTFILES,
        min_font_size=6.5,
        line_height_multiplier=1.5,
    ),
    "ja": PdfLanguageProfile(
        code="ja",
        font_name=PDF_CJK_JP_FONTNAME,
        font_files=PDF_CJK_JP_FONTFILES,
        cache_font_files=PDF_CJK_JP_CACHE_FONTFILES,
        min_font_size=7.5,
        line_height_multiplier=1.3,
    ),
    "ko": PdfLanguageProfile(
        code="ko",
        font_name=PDF_CJK_KR_FONTNAME,
        font_files=PDF_CJK_KR_FONTFILES,
        cache_font_files=PDF_CJK_KR_CACHE_FONTFILES,
        min_font_size=8.0,
        line_height_multiplier=1.4,
    ),
    "ms": PdfLanguageProfile(
        code="ms",
        font_name=PDF_LATIN_FONTNAME,
        font_files=PDF_LATIN_FONTFILES,
        cache_font_files=PDF_LATIN_CACHE_FONTFILES,
        min_font_size=6.0,
        step_down=0.8,
        line_height_multiplier=1.5,
        aggressive_word_wrap=True,
    ),
    "th": PdfLanguageProfile(
        code="th",
        font_name=PDF_THAI_FONTNAME,
        font_files=PDF_THAI_FONTFILES,
        cache_font_files=PDF_THAI_CACHE_FONTFILES,
        min_font_size=7.5,
        line_height_multiplier=1.2,
    ),
    "vi": PdfLanguageProfile(
        code="vi",
        font_name=PDF_LATIN_FONTNAME,
        font_files=PDF_LATIN_FONTFILES,
        cache_font_files=PDF_LATIN_CACHE_FONTFILES,
        min_font_size=7.0,
        line_height_multiplier=1.25,
    ),
}


def preview_sidecar_path(output_path: str) -> Path:
    path = Path(output_path)
    return path.with_name(f"{path.name}.preview.json")


def _utcnow_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _normalize_pdf_block_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", str(text or "")).replace("\r\n", "\n")
    return "\n".join(line.rstrip() for line in normalized.split("\n")).strip()


def _replace_unknown_spacing_block(match: re.Match[str]) -> str:
    source = match.string
    start, end = match.span()
    previous = source[start - 1] if start > 0 else ""
    following = source[end] if end < len(source) else ""
    if previous and following and not previous.isspace() and not following.isspace():
        return "-"
    return " "


def _sanitize_pdf_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", str(text or "")).replace("\r\n", "\n")
    normalized = PDF_CONTROL_CHAR_PATTERN.sub("", normalized)
    normalized = PDF_ZERO_WIDTH_PATTERN.sub("", normalized)
    normalized = PDF_UNKNOWN_SPACING_PATTERN.sub(" ", normalized)
    normalized = PDF_UNKNOWN_GLYPH_BLOCK_PATTERN.sub(_replace_unknown_spacing_block, normalized)
    normalized = PDF_MALFORMED_MULTIPLICATION_PATTERN.sub(" × ", normalized)

    lines = [re.sub(r"[ \t]{2,}", " ", line).strip() for line in normalized.split("\n")]
    return _normalize_pdf_block_text("\n".join(lines))


def _sanitize_pdf_preview_edit_text(text: str) -> tuple[str, bool]:
    normalized_text = _sanitize_pdf_text(text)
    if not normalized_text:
        return "", normalized_text != str(text or "")

    deduped_lines: list[str] = []
    seen_long_lines: set[str] = set()
    changed = normalized_text != str(text or "")

    for line in normalized_text.split("\n"):
        normalized_line = re.sub(r"\s+", "", line)
        if deduped_lines and normalized_line == re.sub(r"\s+", "", deduped_lines[-1]):
            changed = True
            continue
        if len(deduped_lines) >= 2:
            recent_normalized_lines = [re.sub(r"\s+", "", existing_line) for existing_line in deduped_lines[-3:]]
            if len(normalized_line) >= 8 and any(
                normalized_line == f"{left}{right}"
                for index, left in enumerate(recent_normalized_lines)
                for right in recent_normalized_lines[index + 1 :]
            ):
                changed = True
                continue
        if len(normalized_line) >= 8 and normalized_line in seen_long_lines:
            changed = True
            continue
        if len(normalized_line) >= 8:
            similar_existing_line = next(
                (
                    index
                    for index, existing_line in enumerate(deduped_lines)
                    if len(re.sub(r"\s+", "", existing_line)) >= 8
                    and SequenceMatcher(None, normalized_line, re.sub(r"\s+", "", existing_line)).ratio() >= 0.85
                ),
                None,
            )
            if similar_existing_line is not None:
                changed = True
                existing_line = deduped_lines[similar_existing_line]
                if len(normalized_line) > len(re.sub(r"\s+", "", existing_line)):
                    deduped_lines[similar_existing_line] = line
                    seen_long_lines.discard(re.sub(r"\s+", "", existing_line))
                    seen_long_lines.add(normalized_line)
                continue
        deduped_lines.append(line)
        if len(normalized_line) >= 8:
            seen_long_lines.add(normalized_line)

    sanitized = _normalize_pdf_block_text("\n".join(deduped_lines))
    return sanitized, changed or sanitized != normalized_text


def _has_codepoint_in_ranges(text: str, ranges: tuple[tuple[str, str], ...]) -> bool:
    return any(start <= char <= end for char in text for start, end in ranges)


def _select_pdf_font(text: str) -> str:
    if _has_codepoint_in_ranges(text, (("\uac00", "\ud7af"), ("\u1100", "\u11ff"), ("\u3130", "\u318f"))):
        return "korea"
    if _has_codepoint_in_ranges(text, (("\u3040", "\u30ff"), ("\u31f0", "\u31ff"))):
        return "japan"
    if _has_codepoint_in_ranges(
        text,
        (
            ("\u3400", "\u4dbf"),
            ("\u4e00", "\u9fff"),
            ("\uf900", "\ufaff"),
            ("\uff00", "\uffef"),
            ("\u0400", "\u04ff"),
        ),
        ):
        return "china-s"
    return "helv"


def _existing_pdf_font_file(candidates: tuple[str, ...]) -> str | None:
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def _existing_pdf_cached_font_file(candidates: tuple[str, ...]) -> str | None:
    cache_roots = (
        Path.home() / ".cache" / "babeldoc" / "fonts",
        Path("/root/.cache/babeldoc/fonts"),
        Path.cwd() / "tmp" / "babeldoc-cache" / "fonts",
    )
    for root in cache_roots:
        for candidate in candidates:
            font_path = root / candidate
            if font_path.exists():
                return str(font_path)
    return None


def _pdf_language_profile(language: str | None) -> PdfLanguageProfile | None:
    code = PDF_LANGUAGE_ALIASES.get(str(language or "").strip().casefold())
    return PDF_LANGUAGE_PROFILES.get(code or "")


def _pdf_external_profile_render_font(profile: PdfLanguageProfile | None, *, wants_bold: bool = False) -> PdfRenderFont | None:
    if profile is None:
        return None

    route = font_route_for_language(profile.code)
    routed_font_file = route.bold if wants_bold and route and route.bold else route.regular if route else None
    if routed_font_file:
        render_name = f"{profile.font_name}{'-bold' if wants_bold else ''}"
        return PdfRenderFont(
            render_name=render_name,
            render_file=routed_font_file,
            metrics_name=render_name,
            metrics_file=routed_font_file,
        )

    font_file = _existing_pdf_cached_font_file(profile.cache_font_files)
    if font_file is None:
        font_file = _existing_pdf_font_file(profile.font_files)
    if font_file:
        return PdfRenderFont(
            render_name=f"{profile.font_name}{'-bold' if wants_bold else ''}",
            render_file=font_file,
            metrics_name=f"{profile.font_name}{'-bold' if wants_bold else ''}",
            metrics_file=font_file,
        )
    return None


def _pdf_external_cjk_render_font(script_font: str, *, wants_serif: bool) -> PdfRenderFont | None:
    font_name = PDF_CJK_SERIF_FONTNAME if wants_serif else PDF_CJK_SANS_FONTNAME
    font_file = _existing_pdf_cached_font_file(
        PDF_CJK_SERIF_CACHE_FONTFILES if wants_serif else PDF_CJK_SANS_CACHE_FONTFILES
    )
    if font_file is None:
        font_file = _existing_pdf_font_file(PDF_CJK_SERIF_FONTFILES if wants_serif else PDF_CJK_SANS_FONTFILES)
    if font_file is None and wants_serif:
        font_name = PDF_CJK_SANS_FONTNAME
        font_file = _existing_pdf_cached_font_file(PDF_CJK_SANS_CACHE_FONTFILES)
        if font_file is None:
            font_file = _existing_pdf_font_file(PDF_CJK_SANS_FONTFILES)
    if font_file:
        return PdfRenderFont(
            render_name=font_name,
            render_file=font_file,
            metrics_name=font_name,
            metrics_file=font_file,
        )
    return None


def _pdf_external_latin_render_font() -> PdfRenderFont | None:
    font_file = _existing_pdf_cached_font_file(PDF_LATIN_CACHE_FONTFILES)
    if font_file is None:
        font_file = _existing_pdf_font_file(PDF_LATIN_FONTFILES)
    if font_file:
        return PdfRenderFont(
            render_name=PDF_LATIN_FONTNAME,
            render_file=font_file,
            metrics_name=PDF_LATIN_FONTNAME,
            metrics_file=font_file,
        )
    return None


def _pdf_cjk_render_font(script_font: str, *, wants_serif: bool) -> PdfRenderFont:
    external_font = _pdf_external_cjk_render_font(script_font, wants_serif=wants_serif)
    if external_font is not None:
        return external_font
    return PdfRenderFont(render_name=script_font, metrics_name=script_font)


def _pdf_text_needs_embedded_font(text: str) -> bool:
    return any(ord(char) > 255 or char in PDF_EMBEDDED_FONT_CHARS for char in str(text or ""))


def _tokenize_pdf_font_name(font_name: str) -> tuple[str, ...]:
    if not font_name:
        return ()

    expanded = PDF_FONT_SUBSET_PREFIX_PATTERN.sub("", str(font_name or ""))
    expanded = PDF_FONT_CAMEL_CASE_BOUNDARY_PATTERN.sub(" ", expanded)
    tokens = [
        token
        for token in PDF_FONT_SPLIT_PATTERN.split(expanded.casefold())
        if token and token not in PDF_FONT_IGNORED_TOKENS
    ]
    return tuple(tokens)


def _normalize_pdf_font_name(font_name: str) -> str:
    return "".join(_tokenize_pdf_font_name(font_name))


def _pdf_font_name_has_hint(font_name: str, hints: tuple[str, ...]) -> bool:
    normalized_name = _normalize_pdf_font_name(font_name)
    return any(hint in normalized_name for hint in hints)


def _pdf_font_tokens_match(candidate_tokens: tuple[str, ...], preferred_tokens: tuple[str, ...]) -> int:
    score = 0
    for preferred in preferred_tokens:
        if len(preferred) < 3:
            if preferred in candidate_tokens:
                score += 4
            continue
        if any(
            candidate == preferred or candidate.startswith(preferred) or preferred.startswith(candidate)
            for candidate in candidate_tokens
        ):
            score += 10
    return score


def _pdf_font_resource_score(
    resource: PdfPageFontResource,
    preferred_font_name: str,
    *,
    wants_bold: bool,
) -> int:
    preferred_tokens = _tokenize_pdf_font_name(preferred_font_name)
    preferred_normalized = _normalize_pdf_font_name(preferred_font_name)
    score = _pdf_font_tokens_match(resource.tokens, preferred_tokens)
    if preferred_normalized:
        if preferred_normalized == resource.normalized_display_name or preferred_normalized == resource.normalized_resource_name:
            score += 100
        elif preferred_normalized in resource.normalized_display_name or preferred_normalized in resource.normalized_resource_name:
            score += 75
        elif resource.normalized_display_name and resource.normalized_display_name in preferred_normalized:
            score += 60
        elif resource.normalized_resource_name and resource.normalized_resource_name in preferred_normalized:
            score += 60
    if wants_bold == resource.is_bold:
        score += 4
    elif wants_bold and not resource.is_bold:
        score -= 4
    return score


def _font_name_is_serif(font_name: str) -> bool:
    return _pdf_font_name_has_hint(font_name, PDF_SERIF_FONT_HINTS)


def _font_name_is_mono(font_name: str) -> bool:
    return _pdf_font_name_has_hint(font_name, PDF_MONO_FONT_HINTS)


def _font_name_is_symbol(font_name: str) -> bool:
    return _pdf_font_name_has_hint(font_name, PDF_SYMBOL_FONT_HINTS)


def _font_name_is_bold(font_name: str) -> bool:
    return _pdf_font_name_has_hint(font_name, PDF_BOLD_FONT_HINTS)


def _extract_page_font_resources(page: fitz.Page) -> tuple[PdfPageFontResource, ...]:
    resources: list[PdfPageFontResource] = []
    seen_resource_names: set[str] = set()
    document = page.parent

    for font_details in page.get_fonts(full=True):
        font_xref = int(font_details[0])
        resource_name = str(font_details[4] or font_details[3] or "").strip()
        if not resource_name or resource_name in seen_resource_names:
            continue
        seen_resource_names.add(resource_name)

        display_name = str(font_details[3] or resource_name).strip()
        tokens = _tokenize_pdf_font_name(f"{display_name} {resource_name}")
        font_buffer = None
        if document is not None:
            try:
                _, _, _, extracted_buffer = document.extract_font(font_xref)
            except Exception:
                extracted_buffer = b""
            font_buffer = extracted_buffer or None
        resources.append(
            PdfPageFontResource(
                resource_name=resource_name,
                display_name=display_name,
                tokens=tokens,
                normalized_display_name=_normalize_pdf_font_name(display_name),
                normalized_resource_name=_normalize_pdf_font_name(resource_name),
                is_bold=_font_name_is_bold(f"{display_name} {resource_name}"),
                font_buffer=font_buffer,
            )
        )

    return tuple(resources)


def _match_page_font_resource(
    page_fonts: tuple[PdfPageFontResource, ...],
    preferred_font_name: str,
    *,
    wants_bold: bool,
) -> PdfPageFontResource | None:
    if not preferred_font_name:
        return None

    best_match: PdfPageFontResource | None = None
    best_score = 0
    for resource in page_fonts:
        score = _pdf_font_resource_score(resource, preferred_font_name, wants_bold=wants_bold)
        if score > best_score:
            best_match = resource
            best_score = score
    return best_match if best_score >= 20 else None


def _fallback_page_font_resource(
    page_fonts: tuple[PdfPageFontResource, ...],
    *,
    wants_serif: bool,
    wants_mono: bool,
    wants_symbol: bool,
    wants_bold: bool,
) -> PdfPageFontResource | None:
    best_match: PdfPageFontResource | None = None
    best_score = 0

    for resource in page_fonts:
        score = 0
        name = f"{resource.display_name} {resource.resource_name}"
        if wants_symbol and _font_name_is_symbol(name):
            score += 12
        elif wants_mono and _font_name_is_mono(name):
            score += 10
        elif wants_serif and _font_name_is_serif(name):
            score += 10
        elif not wants_serif and _pdf_font_name_has_hint(name, PDF_SANS_FONT_HINTS):
            score += 8
        if wants_bold == resource.is_bold:
            score += 3
        if score > best_score:
            best_match = resource
            best_score = score

    return best_match if best_score > 0 else None


def _resolve_pdf_render_font(
    text: str,
    *,
    preferred_font_name: str = "",
    page_fonts: tuple[PdfPageFontResource, ...] = (),
    prefer_external_cjk_font: bool = False,
    language_profile: PdfLanguageProfile | None = None,
    force_bold: bool = False,
) -> PdfRenderFont:
    wants_serif = _font_name_is_serif(preferred_font_name)
    wants_mono = _font_name_is_mono(preferred_font_name)
    wants_symbol = _font_name_is_symbol(preferred_font_name)
    wants_bold = force_bold or _font_name_is_bold(preferred_font_name)
    script_font = _select_pdf_font(text or preferred_font_name)
    uses_cjk_font = script_font in {"china-s", "japan", "korea"}
    needs_embedded_font = _pdf_text_needs_embedded_font(text)

    if language_profile is not None and language_profile.prefer_external_font and not wants_mono and not wants_symbol:
        external_font = _pdf_external_profile_render_font(language_profile, wants_bold=wants_bold)
        if external_font is not None:
            return external_font

    if ((prefer_external_cjk_font and uses_cjk_font) or needs_embedded_font) and not wants_mono and not wants_symbol:
        external_font = _pdf_external_cjk_render_font(script_font, wants_serif=wants_serif)
        if external_font is not None:
            return external_font
        if needs_embedded_font:
            raise ValueError("Missing embeddable PDF font for edited CJK text. Rebuild the backend image or mount the BabelDOC font cache.")
        return _pdf_cjk_render_font(script_font, wants_serif=wants_serif)

    matched_resource = _match_page_font_resource(page_fonts, preferred_font_name, wants_bold=wants_bold)
    if matched_resource is None:
        matched_resource = _fallback_page_font_resource(
            page_fonts,
            wants_serif=wants_serif,
            wants_mono=wants_mono,
            wants_symbol=wants_symbol,
            wants_bold=wants_bold,
        )

    if matched_resource is not None:
        if uses_cjk_font:
            metrics_font = _pdf_cjk_render_font(script_font, wants_serif=wants_serif)
            metrics_name = metrics_font.metrics_name
            metrics_file = metrics_font.metrics_file
        elif wants_mono:
            metrics_name = "cour"
            metrics_file = None
        elif wants_symbol:
            metrics_name = "symbol"
            metrics_file = None
        elif wants_serif:
            metrics_name = "Times-Roman"
            metrics_file = None
        else:
            metrics_name = "helv"
            metrics_file = None

        render_name = matched_resource.resource_name
        render_buffer = None
        metrics_buffer = None
        if matched_resource.font_buffer:
            render_name = f"dtf_{matched_resource.normalized_resource_name or matched_resource.normalized_display_name or 'font'}"
            render_buffer = matched_resource.font_buffer
            metrics_buffer = matched_resource.font_buffer if not uses_cjk_font else None

        return PdfRenderFont(
            render_name=render_name,
            render_buffer=render_buffer,
            metrics_name=metrics_name,
            metrics_file=metrics_file,
            metrics_buffer=metrics_buffer,
        )

    if uses_cjk_font:
        return _pdf_cjk_render_font(script_font, wants_serif=wants_serif)
    if wants_symbol:
        return PdfRenderFont(render_name="symbol", metrics_name="symbol")
    if wants_mono:
        return PdfRenderFont(render_name="cour", metrics_name="cour")
    if wants_serif:
        return PdfRenderFont(render_name="Times-Roman", metrics_name="Times-Roman")
    return PdfRenderFont(render_name="helv", metrics_name="helv")


@lru_cache(maxsize=16)
def _pdf_metrics_font(font_name: str, font_file: str | None, font_buffer: bytes | None) -> fitz.Font:
    if font_buffer:
        return fitz.Font(fontname=font_name or None, fontbuffer=font_buffer)
    if font_file:
        return fitz.Font(fontname=font_name, fontfile=font_file)
    return fitz.Font(fontname=font_name)


def _pdf_char_uses_cjk_font(char: str) -> bool:
    return _select_pdf_font(char) in {"china-s", "japan", "korea"}


def _pdf_render_font_supports_cjk(font: PdfRenderFont) -> bool:
    cjk_render_names = {
        PDF_CJK_SERIF_FONTNAME,
        PDF_CJK_SANS_FONTNAME,
        PDF_CJK_JP_FONTNAME,
        PDF_CJK_KR_FONTNAME,
        "china-s",
        "japan",
        "korea",
    }
    return any(font.render_name == name or font.render_name.startswith(f"{name}-") for name in cjk_render_names)


def _pdf_should_split_noto_runs(text: str, font: PdfRenderFont) -> bool:
    return bool(font.render_file or font.render_buffer) and any(_pdf_char_uses_cjk_font(char) for char in text) and any(
        not _pdf_char_uses_cjk_font(char) for char in text
    )


def _pdf_latin_render_font(font: PdfRenderFont) -> PdfRenderFont:
    return _pdf_external_latin_render_font() or font


def _pdf_cjk_fallback_render_font(font: PdfRenderFont) -> PdfRenderFont:
    return _pdf_external_cjk_render_font("china-s", wants_serif=False) or font


def _pdf_text_font_runs(text: str, font: PdfRenderFont) -> list[tuple[str, PdfRenderFont]]:
    if not _pdf_should_split_noto_runs(text, font):
        return [(text, font)]

    cjk_font = font if _pdf_render_font_supports_cjk(font) else _pdf_cjk_fallback_render_font(font)
    latin_font = _pdf_latin_render_font(font) if _pdf_render_font_supports_cjk(font) else font
    if cjk_font == latin_font:
        return [(text, font)]

    runs: list[tuple[str, PdfRenderFont]] = []
    current: list[str] = []
    current_font = cjk_font if _pdf_char_uses_cjk_font(text[0]) else latin_font
    for char in text:
        next_font = cjk_font if _pdf_char_uses_cjk_font(char) else latin_font
        if next_font != current_font:
            runs.append(("".join(current), current_font))
            current = []
            current_font = next_font
        current.append(char)
    if current:
        runs.append(("".join(current), current_font))
    return runs


def _rect_iou(left: fitz.Rect, right: fitz.Rect) -> float:
    intersection_width = max(0.0, min(left.x1, right.x1) - max(left.x0, right.x0))
    intersection_height = max(0.0, min(left.y1, right.y1) - max(left.y0, right.y0))
    intersection_area = intersection_width * intersection_height
    if intersection_area <= 0:
        return 0.0

    left_area = max(0.0, left.width) * max(0.0, left.height)
    right_area = max(0.0, right.width) * max(0.0, right.height)
    union_area = left_area + right_area - intersection_area
    return intersection_area / union_area if union_area > 0 else 0.0


def _rect_overlap_ratio(left: fitz.Rect, right: fitz.Rect) -> float:
    intersection_width = max(0.0, min(left.x1, right.x1) - max(left.x0, right.x0))
    intersection_height = max(0.0, min(left.y1, right.y1) - max(left.y0, right.y0))
    intersection_area = intersection_width * intersection_height
    left_area = max(0.0, left.width) * max(0.0, left.height)
    return intersection_area / left_area if left_area > 0 else 0.0


def _rect_vertical_gap(left: fitz.Rect, right: fitz.Rect) -> float:
    if left.y1 < right.y0:
        return right.y0 - left.y1
    if right.y1 < left.y0:
        return left.y0 - right.y1
    return 0.0


def _rect_horizontal_gap(left: fitz.Rect, right: fitz.Rect) -> float:
    if left.x1 < right.x0:
        return right.x0 - left.x1
    if right.x1 < left.x0:
        return left.x0 - right.x1
    return 0.0


def _horizontal_overlap_ratio(left: fitz.Rect, right: fitz.Rect) -> float:
    overlap = max(0.0, min(left.x1, right.x1) - max(left.x0, right.x0))
    smallest_width = max(min(left.width, right.width), 1.0)
    return overlap / smallest_width


def _pdf_fragment_column_domain(fragment: PdfPreviewFragment, page_width: float) -> str | None:
    if page_width <= 0:
        return None
    midline = page_width / 2
    center_x = (fragment.rect.x0 + fragment.rect.x1) / 2
    if center_x < midline and fragment.rect.x1 <= midline + PDF_BLOCK_MERGE_HORIZONTAL_GAP:
        return "left"
    if center_x > midline and fragment.rect.x0 >= midline - PDF_BLOCK_MERGE_HORIZONTAL_GAP:
        return "right"
    return None


def _should_merge_pdf_fragments(
    left: PdfPreviewFragment,
    right: PdfPreviewFragment,
    *,
    page_width: float = 0.0,
    isolate_page_columns: bool = False,
) -> bool:
    if isolate_page_columns:
        left_domain = _pdf_fragment_column_domain(left, page_width)
        right_domain = _pdf_fragment_column_domain(right, page_width)
        if left_domain is not None and right_domain is not None and left_domain != right_domain:
            return False

    if _rect_iou(left.rect, right.rect) > PDF_BLOCK_MERGE_IOU_THRESHOLD:
        return True

    same_reading_lane = (
        _horizontal_overlap_ratio(left.rect, right.rect) >= PDF_BLOCK_MERGE_MIN_HORIZONTAL_OVERLAP_RATIO
        or _rect_horizontal_gap(left.rect, right.rect) <= PDF_BLOCK_MERGE_HORIZONTAL_GAP
    )
    return _rect_vertical_gap(left.rect, right.rect) <= PDF_BLOCK_MERGE_VERTICAL_DISTANCE and same_reading_lane


def _pdf_fragment_sort_key(fragment: PdfPreviewFragment) -> tuple[float, float]:
    return (round(fragment.rect.y0, 2), round(fragment.rect.x0, 2))


def _pdf_fragment_center_y(fragment: PdfPreviewFragment) -> float:
    return (fragment.rect.y0 + fragment.rect.y1) / 2


def _pdf_visual_line_center_tolerance(fragments: list[PdfPreviewFragment]) -> float:
    heights = [fragment.rect.height for fragment in fragments if fragment.rect.height > 0]
    height_tolerance = _median(heights) * PDF_VISUAL_LINE_HEIGHT_TOLERANCE_RATIO
    return min(
        max(height_tolerance, PDF_VISUAL_LINE_MIN_CENTER_TOLERANCE),
        PDF_VISUAL_LINE_MAX_CENTER_TOLERANCE,
    )


def _group_pdf_fragments_by_visual_line(fragments: list[PdfPreviewFragment]) -> list[list[PdfPreviewFragment]]:
    if not fragments:
        return []

    tolerance = _pdf_visual_line_center_tolerance(fragments)
    groups: list[list[PdfPreviewFragment]] = []
    group_centers: list[float] = []

    for fragment in sorted(fragments, key=lambda item: (_pdf_fragment_center_y(item), item.rect.x0)):
        center_y = _pdf_fragment_center_y(fragment)
        if groups and abs(center_y - group_centers[-1]) <= tolerance:
            groups[-1].append(fragment)
            group_centers[-1] = sum(_pdf_fragment_center_y(item) for item in groups[-1]) / len(groups[-1])
            continue
        groups.append([fragment])
        group_centers.append(center_y)

    return [sorted(group, key=lambda item: (round(item.rect.x0, 2), round(item.rect.y0, 2))) for group in groups]


def _sort_pdf_fragments_by_visual_lines(fragments: list[PdfPreviewFragment]) -> list[PdfPreviewFragment]:
    return [fragment for group in _group_pdf_fragments_by_visual_line(fragments) for fragment in group]


def _merge_pdf_fragment_rects(fragments: list[PdfPreviewFragment]) -> fitz.Rect:
    return fitz.Rect(
        min(fragment.rect.x0 for fragment in fragments),
        min(fragment.rect.y0 for fragment in fragments),
        max(fragment.rect.x1 for fragment in fragments),
        max(fragment.rect.y1 for fragment in fragments),
    )


def _merge_pdf_fragment_texts(fragments: list[PdfPreviewFragment]) -> str:
    lines = [
        " ".join(fragment.text.strip() for fragment in group if fragment.text.strip()).strip()
        for group in _group_pdf_fragments_by_visual_line(fragments)
    ]
    return _normalize_pdf_block_text("\n".join(line for line in lines if line))


def _merge_pdf_fragments(fragments: list[PdfPreviewFragment]) -> PdfPreviewFragment:
    return PdfPreviewFragment(
        rect=_merge_pdf_fragment_rects(fragments),
        text=_merge_pdf_fragment_texts(fragments),
        font_names=[font_name for fragment in fragments for font_name in fragment.font_names],
        font_sizes=[font_size for fragment in fragments for font_size in fragment.font_sizes],
        rotations=[rotation for fragment in fragments for rotation in fragment.rotations],
    )


def _pdf_text_starts_new_list_item(text: str) -> bool:
    return bool(PDF_BULLET_TEXT_PATTERN.match(str(text or "").strip()))


def _pdf_technical_assignment_key(text: str) -> str | None:
    match = PDF_TECHNICAL_ASSIGNMENT_PATTERN.match(str(text or ""))
    if match is None:
        return None
    return match.group(1).casefold()


def _has_repeating_pdf_technical_assignments(fragments: list[PdfPreviewFragment]) -> bool:
    keys = [_pdf_technical_assignment_key(fragment.text) for fragment in fragments]
    keys = [key for key in keys if key is not None]
    return len(keys) >= 2 and len(set(keys)) < len(keys)


def _pdf_fragment_group_vertical_step(fragments: list[PdfPreviewFragment]) -> float:
    centers = [(fragment.rect.y0 + fragment.rect.y1) / 2 for fragment in fragments]
    steps = [
        centers[index + 1] - centers[index]
        for index in range(len(centers) - 1)
        if centers[index + 1] - centers[index] > 1.0
    ]
    heights = [fragment.rect.height for fragment in fragments if fragment.rect.height > 0]
    return max(_median(steps), _median(heights), PDF_MIN_REDRAW_FONT_SIZE + 1.0)


def _cluster_pdf_fragment_groups(
    fragments: list[PdfPreviewFragment],
    *,
    page_width: float = 0.0,
    page_index: int | None = None,
) -> list[list[PdfPreviewFragment]]:
    remaining = _sort_pdf_fragments_by_visual_lines(fragments)
    groups: list[list[PdfPreviewFragment]] = []
    isolate_page_columns = False

    while remaining:
        cluster = [remaining.pop(0)]
        cluster_changed = True
        while cluster_changed:
            cluster_changed = False
            unmatched: list[PdfPreviewFragment] = []
            for candidate in remaining:
                if any(
                    _should_merge_pdf_fragments(
                        member,
                        candidate,
                        page_width=page_width,
                        isolate_page_columns=isolate_page_columns,
                    )
                    for member in cluster
                ):
                    cluster.append(candidate)
                    cluster_changed = True
                else:
                    unmatched.append(candidate)
            remaining = unmatched

        groups.append(_sort_pdf_fragments_by_visual_lines(cluster))

    return sorted(groups, key=lambda group: _pdf_fragment_sort_key(group[0]))


def _pdf_fragment_starts_new_block(
    current_group: list[PdfPreviewFragment],
    fragment: PdfPreviewFragment,
    *,
    max_vertical_gap: float,
) -> bool:
    if not current_group:
        return False

    if _pdf_text_starts_new_list_item(fragment.text):
        return True

    previous = current_group[-1]
    if _rect_vertical_gap(previous.rect, fragment.rect) > max_vertical_gap:
        return True

    group_rect = _merge_pdf_fragment_rects(current_group)
    if len(current_group) >= PDF_BLOCK_GROUP_MAX_LINES:
        return True
    if (group_rect | fragment.rect).height > PDF_BLOCK_GROUP_MAX_HEIGHT:
        return True

    group_left = min(item.rect.x0 for item in current_group)
    previous_indent = previous.rect.x0 - group_left
    current_indent = fragment.rect.x0 - group_left
    return previous_indent >= 8.0 and current_indent <= 2.0


def _split_pdf_fragment_group(
    fragments: list[PdfPreviewFragment],
    *,
    page_width: float = 0.0,
    page_index: int | None = None,
) -> list[PdfPreviewFragment]:
    line_fragments = _merge_line_fragments_for_layout(fragments)
    if len(line_fragments) <= 1:
        return line_fragments
    if _has_repeating_pdf_technical_assignments(line_fragments):
        return line_fragments

    vertical_step = _pdf_fragment_group_vertical_step(line_fragments)
    max_vertical_gap = max(PDF_BLOCK_MERGE_VERTICAL_DISTANCE, vertical_step * PDF_BLOCK_GROUP_PARAGRAPH_GAP_RATIO)
    blocks: list[PdfPreviewFragment] = []
    current_group: list[PdfPreviewFragment] = []

    for fragment in line_fragments:
        if _pdf_fragment_starts_new_block(current_group, fragment, max_vertical_gap=max_vertical_gap):
            blocks.append(_merge_pdf_fragments(current_group))
            current_group = []
        current_group.append(fragment)

    if current_group:
        blocks.append(_merge_pdf_fragments(current_group))

    return blocks


def _cluster_pdf_fragments(
    fragments: list[PdfPreviewFragment],
    *,
    page_width: float = 0.0,
    page_index: int | None = None,
) -> list[PdfPreviewFragment]:
    merged_fragments: list[PdfPreviewFragment] = []
    for group in _cluster_pdf_fragment_groups(fragments, page_width=page_width, page_index=page_index):
        merged_fragments.extend(_split_pdf_fragment_group(group, page_width=page_width, page_index=page_index))
    return _sort_pdf_fragments_by_visual_lines(merged_fragments)


def _extract_pdf_text_fragments(page: fitz.Page) -> list[PdfPreviewFragment]:
    text_dict = page.get_text("dict")
    fragments: list[PdfPreviewFragment] = []

    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue

        for line in block.get("lines", []):
            line_parts: list[str] = []
            font_names: list[str] = []
            font_sizes: list[float] = []
            rotations: list[int] = []
            for span in line.get("spans", []):
                span_text = span.get("text", "")
                if not span_text.strip():
                    continue
                span_size = span.get("size")
                if isinstance(span_size, (int, float)) and float(span_size) < PDF_MIN_VISIBLE_TEXT_SPAN_SIZE:
                    continue
                line_parts.append(span_text)
                font_name = str(span.get("font", "")).strip()
                if font_name:
                    font_names.append(font_name)
                if isinstance(span_size, (int, float)):
                    font_sizes.append(float(span_size))
                rotation = _pdf_rotation_from_direction(span.get("dir") or line.get("dir"))
                if rotation is not None:
                    rotations.append(rotation)
            line_text = _sanitize_pdf_text("".join(line_parts))
            if not line_text:
                continue

            bbox = fitz.Rect(line.get("bbox") or block.get("bbox", (0, 0, 0, 0)))
            fragments.append(
                PdfPreviewFragment(
                    rect=bbox,
                    text=line_text,
                    font_names=font_names,
                    font_sizes=font_sizes,
                    rotations=rotations,
                )
            )

    return fragments


def _extract_pdf_words(page: fitz.Page | None) -> list[PdfPreviewWord]:
    if page is None:
        return []

    words: list[PdfPreviewWord] = []
    for raw_word in page.get_text("words"):
        text = _sanitize_pdf_text(raw_word[4])
        if not text:
            continue

        words.append(
            PdfPreviewWord(
                rect=fitz.Rect(raw_word[:4]),
                text=text,
                block_index=int(raw_word[5]),
                line_index=int(raw_word[6]),
                word_index=int(raw_word[7]),
            )
        )

    return words


def _extract_page_text_in_rect(
    page: fitz.Page | None,
    rect: fitz.Rect,
    *,
    textpage: fitz.TextPage | None = None,
) -> str:
    if page is None:
        return ""
    return _sanitize_pdf_preview_edit_text(page.get_textbox(rect, textpage=textpage))[0]


def _extract_clipped_words_text_in_rect(
    page: fitz.Page | None,
    rect: fitz.Rect,
    *,
    textpage: fitz.TextPage | None = None,
    padding: float = 1.0,
) -> str:
    if page is None or textpage is None:
        return ""

    clip = fitz.Rect(rect.x0 - padding, rect.y0 - padding, rect.x1 + padding, rect.y1 + padding)
    lines: list[list[str]] = []
    current_line_key: tuple[int, int] | None = None
    current_words: list[str] = []

    for raw_word in page.get_text("words", clip=clip, textpage=textpage):
        line_key = (int(raw_word[5]), int(raw_word[6]))
        if current_line_key is not None and line_key != current_line_key:
            if current_words:
                lines.append(current_words)
            current_words = []

        text = _sanitize_pdf_text(raw_word[4])
        if text:
            current_words.append(text)
        current_line_key = line_key

    if current_words:
        lines.append(current_words)

    extracted_text = _normalize_pdf_block_text("\n".join(" ".join(line).strip() for line in lines if line))
    sanitized_text = _sanitize_pdf_preview_edit_text(extracted_text)[0]
    return sanitized_text or _extract_page_text_in_rect(page, rect, textpage=textpage)


def _word_belongs_to_rect(word: PdfPreviewWord, rect: fitz.Rect) -> bool:
    center_x = (word.rect.x0 + word.rect.x1) / 2
    center_y = (word.rect.y0 + word.rect.y1) / 2
    if rect.x0 <= center_x <= rect.x1 and rect.y0 <= center_y <= rect.y1:
        return True
    return _rect_overlap_ratio(word.rect, rect) >= 0.5 or _rect_overlap_ratio(rect, word.rect) >= 0.5


def _fragment_belongs_to_rect(fragment: PdfPreviewFragment, rect: fitz.Rect) -> bool:
    center_x = (fragment.rect.x0 + fragment.rect.x1) / 2
    center_y = (fragment.rect.y0 + fragment.rect.y1) / 2
    if rect.x0 <= center_x <= rect.x1 and rect.y0 <= center_y <= rect.y1:
        return True
    return _rect_overlap_ratio(fragment.rect, rect) >= 0.5 or _rect_overlap_ratio(rect, fragment.rect) >= 0.5


def _extract_words_text_in_rect(words: list[PdfPreviewWord], rect: fitz.Rect) -> str:
    selected_words: list[PdfPreviewWord] = []
    seen_words: set[tuple[float, float, float, float, str]] = set()
    for word in words:
        if not _word_belongs_to_rect(word, rect):
            continue
        word_key = (
            round(word.rect.x0, 2),
            round(word.rect.y0, 2),
            round(word.rect.x1, 2),
            round(word.rect.y1, 2),
            word.text,
        )
        if word_key in seen_words:
            continue
        seen_words.add(word_key)
        selected_words.append(word)
    if not selected_words:
        return ""

    ordered_words = sorted(
        selected_words,
        key=lambda word: (
            round(word.rect.y0, 2),
            round(word.rect.x0, 2),
            word.block_index,
            word.line_index,
            word.word_index,
        ),
    )

    lines: list[list[str]] = []
    current_line: list[str] = []
    current_top: float | None = None
    current_block: int | None = None
    current_line_index: int | None = None

    for word in ordered_words:
        starts_new_line = (
            current_top is not None
            and (
                abs(word.rect.y0 - current_top) > 2.0
                or word.block_index != current_block
                or word.line_index != current_line_index
            )
        )
        if starts_new_line and current_line:
            lines.append(current_line)
            current_line = []

        current_line.append(word.text)
        current_top = word.rect.y0
        current_block = word.block_index
        current_line_index = word.line_index

    if current_line:
        lines.append(current_line)

    extracted_text = _normalize_pdf_block_text("\n".join(" ".join(line).strip() for line in lines if line))
    return _sanitize_pdf_preview_edit_text(extracted_text)[0]


def _extract_fragment_text_in_rect(fragments: list[PdfPreviewFragment], rect: fitz.Rect) -> str:
    selected_fragments = [
        fragment
        for fragment in fragments
        if _fragment_belongs_to_rect(fragment, rect)
    ]
    if not selected_fragments:
        return ""
    return _sanitize_pdf_preview_edit_text(_merge_pdf_fragment_texts(selected_fragments))[0]


def _average_font_size_in_rect(fragments: list[PdfPreviewFragment], rect: fitz.Rect) -> float:
    font_sizes = [
        font_size
        for fragment in fragments
        if _rect_overlap_ratio(fragment.rect, rect) >= 0.5 or _rect_overlap_ratio(rect, fragment.rect) >= 0.5
        for font_size in fragment.font_sizes
    ]
    return round(sum(font_sizes) / len(font_sizes), 2) if font_sizes else 11.0


def _dominant_font_name_in_rect(fragments: list[PdfPreviewFragment], rect: fitz.Rect, *, fallback_text: str = "") -> str:
    font_names = [
        font_name
        for fragment in fragments
        if _rect_overlap_ratio(fragment.rect, rect) >= 0.5 or _rect_overlap_ratio(rect, fragment.rect) >= 0.5
        for font_name in fragment.font_names
        if font_name
    ]
    if font_names:
        return Counter(font_names).most_common(1)[0][0]
    return _select_pdf_font(fallback_text)


def _pdf_rotation_from_direction(direction: object) -> int | None:
    if not isinstance(direction, (list, tuple)) or len(direction) < 2:
        return None
    dx, dy = direction[:2]
    if not isinstance(dx, (int, float)) or not isinstance(dy, (int, float)):
        return None
    if abs(float(dx)) > PDF_ROTATION_VECTOR_TOLERANCE:
        return None
    if float(dy) <= -1 + PDF_ROTATION_VECTOR_TOLERANCE:
        return 90
    if float(dy) >= 1 - PDF_ROTATION_VECTOR_TOLERANCE:
        return 270
    return None


def _pdf_alignment_for_rect(page_width: float, rect: fitz.Rect) -> str | None:
    rect_mid = (rect.x0 + rect.x1) / 2
    return sniff_alignment(page_width, rect_mid, tolerance=PDF_CENTER_ALIGNMENT_TOLERANCE)


def _pdf_item_metadata(rect: fitz.Rect, page_width: float, fragments: list[PdfPreviewFragment]) -> dict[str, object]:
    selected_fragments = [fragment for fragment in fragments if _fragment_belongs_to_rect(fragment, rect)]
    metadata: dict[str, object] = {"layout_status": "ok"}
    alignment = _pdf_alignment_for_rect(page_width, rect)
    if alignment is not None:
        metadata["alignment"] = alignment
    if any(fragment.is_bold for fragment in selected_fragments):
        metadata["font_style"] = "BOLD"
    rotations = [fragment.dominant_rotation for fragment in selected_fragments]
    rotations = [rotation for rotation in rotations if rotation is not None]
    if rotations:
        metadata["rotation"] = Counter(rotations).most_common(1)[0][0]
    return metadata


def _median(values: list[float]) -> float:
    if not values:
        return 0.0

    ordered = sorted(values)
    middle = len(ordered) // 2
    if len(ordered) % 2:
        return ordered[middle]
    return (ordered[middle - 1] + ordered[middle]) / 2


def _line_fragments_in_rect(fragments: list[PdfPreviewFragment], rect: fitz.Rect) -> list[PdfPreviewFragment]:
    return _sort_pdf_fragments_by_visual_lines([fragment for fragment in fragments if _fragment_belongs_to_rect(fragment, rect)])


def _merge_line_fragments_for_layout(fragments: list[PdfPreviewFragment]) -> list[PdfPreviewFragment]:
    return [_merge_pdf_fragments(group) for group in _group_pdf_fragments_by_visual_line(fragments)]


def _build_pdf_line_rects(
    rect: fitz.Rect,
    line_fragments: list[PdfPreviewFragment],
    font_size: float,
    *,
    full_width: bool = False,
    allow_extra_lines: bool = False,
    line_height_multiplier: float = 1.12,
) -> tuple[list[fitz.Rect], int]:
    if not line_fragments:
        return [], 0

    layout_fragments = _merge_line_fragments_for_layout(line_fragments)
    min_x0 = max(rect.x0, min(fragment.rect.x0 for fragment in layout_fragments) - 1.5)
    max_x1 = min(rect.x1, max(fragment.rect.x1 for fragment in layout_fragments) + 1.5)
    fragment_heights = [fragment.rect.height for fragment in layout_fragments if fragment.rect.height > 0]
    fragment_centers = [(fragment.rect.y0 + fragment.rect.y1) / 2 for fragment in layout_fragments]
    line_height = max(_median(fragment_heights), font_size * line_height_multiplier, PDF_MIN_REDRAW_FONT_SIZE + 1.0)
    vertical_steps = [
        fragment_centers[index + 1] - fragment_centers[index]
        for index in range(len(fragment_centers) - 1)
        if fragment_centers[index + 1] - fragment_centers[index] > 1.0
    ]
    vertical_step = max(_median(vertical_steps), font_size * 1.05, PDF_MIN_REDRAW_FONT_SIZE + 0.5)

    line_rects: list[fitz.Rect] = []
    for fragment, center_y in zip(layout_fragments, fragment_centers, strict=False):
        y0 = max(rect.y0, center_y - line_height / 2)
        y1 = min(rect.y1, center_y + line_height / 2)
        if y1 - y0 < font_size:
            y1 = min(rect.y1, y0 + max(line_height, font_size * 1.1))
        if full_width:
            x0 = min_x0
            x1 = max_x1
        else:
            x0 = max(rect.x0, fragment.rect.x0 - 1.5)
            x1 = min(rect.x1, fragment.rect.x1 + 1.5)
        line_rects.append(fitz.Rect(x0, y0, x1, y1))

    base_count = len(line_rects)
    if not allow_extra_lines:
        return line_rects, base_count

    next_center_y = fragment_centers[-1] + vertical_step
    last_line_rect = line_rects[-1]
    while next_center_y + line_height / 2 <= rect.y1 + 0.5:
        y0 = max(rect.y0, next_center_y - line_height / 2)
        y1 = min(rect.y1, next_center_y + line_height / 2)
        line_rects.append(fitz.Rect(last_line_rect.x0, y0, last_line_rect.x1, y1))
        next_center_y += vertical_step

    return line_rects, base_count


def _tokenize_pdf_wrap_text(text: str) -> list[str]:
    tokens: list[str] = []
    for token in PDF_WRAP_TOKEN_PATTERN.findall(text):
        if token.isspace():
            if tokens and not tokens[-1].isspace():
                tokens.append(" ")
            continue
        tokens.append(token)

    while tokens and tokens[-1].isspace():
        tokens.pop()
    return tokens


def _pdf_text_width(text: str, font: PdfRenderFont, font_size: float) -> float:
    if not text:
        return 0.0
    if _pdf_should_split_noto_runs(text, font):
        runs = _pdf_text_font_runs(text, font)
        if len(runs) > 1 or runs[0][1] != font:
            return sum(_pdf_text_width(run_text, run_font, font_size) for run_text, run_font in runs)
    return _pdf_metrics_font(font.metrics_name, font.metrics_file, font.metrics_buffer).text_length(
        text,
        fontsize=font_size,
    )


def _pdf_text_width_single_font(text: str, font: PdfRenderFont, font_size: float) -> float:
    if not text:
        return 0.0
    return _pdf_metrics_font(font.metrics_name, font.metrics_file, font.metrics_buffer).text_length(
        text,
        fontsize=font_size,
    )


def _split_pdf_wrap_token(token: str, max_width: float, font: PdfRenderFont, font_size: float) -> tuple[str, str]:
    if not token:
        return "", ""

    prefix = ""
    for index, char in enumerate(token, start=1):
        candidate = prefix + char
        if _pdf_text_width(candidate, font, font_size) <= max_width:
            prefix = candidate
            continue
        if prefix:
            return prefix, token[index - 1 :]
        return "", token

    return prefix, ""


def _wrap_pdf_text_to_widths(
    text: str,
    widths: list[float],
    font: PdfRenderFont,
    font_size: float,
) -> list[str] | None:
    paragraphs = text.split("\n")
    lines: list[str] = []
    line_index = 0

    for paragraph in paragraphs:
        if paragraph == "":
            if line_index >= len(widths):
                return None
            lines.append("")
            line_index += 1
            continue

        tokens = _tokenize_pdf_wrap_text(paragraph)
        current = ""

        while tokens:
            if line_index >= len(widths):
                return None

            token = tokens.pop(0)
            width = max(widths[line_index] - 1.0, 1.0)
            candidate = f"{current}{token}"

            if candidate and _pdf_text_width(candidate, font, font_size) <= width:
                current = candidate
                continue

            if current:
                lines.append(current.rstrip())
                line_index += 1
                current = ""
                stripped_token = token.lstrip()
                if stripped_token:
                    tokens.insert(0, stripped_token)
                continue

            split_token, remainder = _split_pdf_wrap_token(token.lstrip(), width, font, font_size)
            if not split_token:
                return None
            lines.append(split_token.rstrip())
            line_index += 1
            if remainder:
                tokens.insert(0, remainder)

        if current:
            if line_index >= len(widths):
                return None
            lines.append(current.rstrip())
            line_index += 1

    return lines


def _build_pdf_text_blocks(
    fragments: list[PdfPreviewFragment],
    *,
    page_index: int,
    source_fragments: list[PdfPreviewFragment],
    target_fragments: list[PdfPreviewFragment],
    page_width: float,
) -> list[dict[str, object]]:
    merged_blocks = _cluster_pdf_fragments(fragments, page_width=page_width, page_index=page_index)
    return [
        {
            "type": "text",
            "rect": [round(block.rect.x0, 2), round(block.rect.y0, 2), round(block.rect.x1, 2), round(block.rect.y1, 2)],
            "font_name": block.dominant_font,
            "font_size_original": block.average_font_size,
            "font_size_current": block.average_font_size,
            "src_text": _extract_fragment_text_in_rect(source_fragments, block.rect),
            "tgt_text": _extract_fragment_text_in_rect(target_fragments, block.rect),
            **_pdf_item_metadata(block.rect, page_width, fragments),
        }
        for block in merged_blocks
    ]


def _dedupe_sorted_positions(values: list[float], *, tolerance: float = 1.0) -> list[float]:
    deduped: list[float] = []
    for value in sorted(values):
        if deduped and abs(value - deduped[-1]) <= tolerance:
            deduped[-1] = max(deduped[-1], value)
        else:
            deduped.append(value)
    return deduped


def _table_grid_positions(table) -> tuple[list[float], list[float]]:
    x_positions = [float(table.bbox[0]), float(table.bbox[2])]
    y_positions = [float(table.bbox[1]), float(table.bbox[3])]

    for row in table.rows:
        for cell in row.cells:
            if cell is None:
                continue
            x0, y0, x1, y1 = cell
            x_positions.extend([float(x0), float(x1)])
            y_positions.extend([float(y0), float(y1)])

    return _dedupe_sorted_positions(x_positions), _dedupe_sorted_positions(y_positions)


def _grid_index_for_position(positions: list[float], value: float, *, tolerance: float = 1.0) -> int:
    for index, position in enumerate(positions):
        if abs(position - value) <= tolerance:
            return index
    return min(range(len(positions)), key=lambda index: abs(positions[index] - value))


def _fragment_belongs_to_table(fragment: PdfPreviewFragment, table_rects: list[fitz.Rect]) -> bool:
    center_x = (fragment.rect.x0 + fragment.rect.x1) / 2
    center_y = (fragment.rect.y0 + fragment.rect.y1) / 2
    for table_rect in table_rects:
        if table_rect.x0 <= center_x <= table_rect.x1 and table_rect.y0 <= center_y <= table_rect.y1:
            return True
        if _rect_overlap_ratio(fragment.rect, table_rect) >= 0.5:
            return True
    return False


def _extract_pdf_table_blocks(
    editable_page: fitz.Page | None,
    *,
    page_index: int,
    page_width: float,
    source_page: fitz.Page | None,
    source_textpage: fitz.TextPage | None,
    target_words: list[PdfPreviewWord],
    source_fragments: list[PdfPreviewFragment],
    target_fragments: list[PdfPreviewFragment],
) -> list[dict[str, object]]:
    if editable_page is None:
        return []

    table_blocks: list[dict[str, object]] = []
    for table_index, table in enumerate(editable_page.find_tables().tables, start=1):
        x_positions, y_positions = _table_grid_positions(table)
        rows_count = max(len(y_positions) - 1, 0)
        cols_count = max(len(x_positions) - 1, 0)
        if rows_count < PDF_TABLE_MIN_ROWS or cols_count < PDF_TABLE_MIN_COLS:
            continue

        cells: list[dict[str, object]] = []
        seen_rects: set[tuple[float, float, float, float]] = set()
        font_fragments = target_fragments or source_fragments
        for row in table.rows:
            for cell in row.cells:
                if cell is None:
                    continue

                rect = fitz.Rect(cell)
                rect_key = (round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2))
                if rect_key in seen_rects:
                    continue
                seen_rects.add(rect_key)

                row_start = _grid_index_for_position(y_positions, rect.y0)
                row_end = _grid_index_for_position(y_positions, rect.y1)
                col_start = _grid_index_for_position(x_positions, rect.x0)
                col_end = _grid_index_for_position(x_positions, rect.x1)
                font_size = _average_font_size_in_rect(font_fragments, rect)
                source_text = _extract_clipped_words_text_in_rect(
                    source_page,
                    rect,
                    textpage=source_textpage,
                )
                target_text = _extract_words_text_in_rect(target_words, rect)
                font_name = _dominant_font_name_in_rect(
                    font_fragments,
                    rect,
                    fallback_text=target_text or source_text,
                )

                cells.append(
                    {
                        "cell_id": f"p{page_index + 1}_t{table_index}_r{row_start + 1}_c{col_start + 1}",
                        "row_index": row_start + 1,
                        "col_index": col_start + 1,
                        "row_span": max(row_end - row_start, 1),
                        "col_span": max(col_end - col_start, 1),
                        "rect": [round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2)],
                        "font_name": font_name,
                        "font_size_original": font_size,
                        "font_size_current": font_size,
                        "src_text": source_text,
                        "tgt_text": target_text,
                        **_pdf_item_metadata(rect, page_width, font_fragments),
                    }
                )

        table_cell_density = len(cells) / max(rows_count * cols_count, 1)
        if len(cells) < PDF_TABLE_MIN_CELLS or table_cell_density < PDF_TABLE_MIN_CELL_DENSITY:
            continue

        table_blocks.append(
            {
                "type": "table",
                "block_id": f"p{page_index + 1}_t{table_index}",
                "table_rect": [
                    round(float(table.bbox[0]), 2),
                    round(float(table.bbox[1]), 2),
                    round(float(table.bbox[2]), 2),
                    round(float(table.bbox[3]), 2),
                ],
                "rows_count": rows_count,
                "cols_count": cols_count,
                "cells": cells,
            }
        )

    return table_blocks


def _page_item_sort_key(item: dict[str, object]) -> tuple[float, float]:
    rect = item.get("table_rect") if item.get("type") == "table" else item.get("rect")
    if not isinstance(rect, list) or len(rect) != 4:
        return (0.0, 0.0)
    return (float(rect[1]), float(rect[0]))


def _docx_paragraph_texts(path: Path) -> list[str]:
    if not path.exists():
        return []
    document = Document(path)
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
    return texts


def _build_pdf_pages(source_path: Path, output_path: Path) -> list[dict[str, object]]:
    source_document = fitz.open(source_path) if source_path.exists() else None
    output_document = fitz.open(output_path)
    try:
        page_count = max(source_document.page_count if source_document is not None else 0, output_document.page_count, 1)
        pages: list[dict[str, object]] = []

        for page_index in range(page_count):
            source_page = (
                source_document.load_page(page_index)
                if source_document is not None and page_index < source_document.page_count
                else None
            )
            output_page = output_document.load_page(page_index) if page_index < output_document.page_count else None
            geometry_page = output_page or source_page
            source_fragments = _extract_pdf_text_fragments(source_page) if source_page is not None else []
            target_fragments = _extract_pdf_text_fragments(output_page) if output_page is not None else []
            source_textpage = source_page.get_textpage() if source_page is not None else None
            target_words = _extract_pdf_words(output_page)
            geometry_fragments = source_fragments or target_fragments
            page_width = round(geometry_page.rect.width, 2) if geometry_page is not None else 0.0
            page_height = round(geometry_page.rect.height, 2) if geometry_page is not None else 0.0
            table_blocks = _extract_pdf_table_blocks(
                output_page or source_page,
                page_index=page_index,
                page_width=page_width,
                source_page=source_page,
                source_textpage=source_textpage,
                target_words=target_words,
                source_fragments=source_fragments,
                target_fragments=target_fragments,
            )
            table_rects = [fitz.Rect(block["table_rect"]) for block in table_blocks]
            filtered_source_fragments = [fragment for fragment in source_fragments if not _fragment_belongs_to_table(fragment, table_rects)]
            filtered_target_fragments = [fragment for fragment in target_fragments if not _fragment_belongs_to_table(fragment, table_rects)]
            text_blocks = _build_pdf_text_blocks(
                [fragment for fragment in geometry_fragments if not _fragment_belongs_to_table(fragment, table_rects)],
                page_index=page_index,
                source_fragments=filtered_source_fragments,
                target_fragments=filtered_target_fragments,
                page_width=page_width,
            )
            page_items = sorted([*text_blocks, *table_blocks], key=_page_item_sort_key)

            blocks: list[dict[str, object]] = []

            text_block_index = 0
            for block in page_items:
                if block.get("type") == "table":
                    blocks.append(block)
                    continue

                text_block_index += 1
                blocks.append(
                    {
                        "type": "text",
                        "block_id": f"p{page_index + 1}_b{text_block_index}",
                        "rect": block["rect"],
                        "font_name": str(block.get("font_name", "")),
                        "font_size_original": float(block.get("font_size_original", 12.0) or 12.0),
                        "font_size_current": float(block.get("font_size_current", 12.0) or 12.0),
                        "src_text": str(block.get("src_text", "")),
                        "tgt_text": str(block.get("tgt_text", "")),
                        "alignment": block.get("alignment"),
                        "font_style": block.get("font_style"),
                        "rotation": block.get("rotation"),
                        "layout_status": block.get("layout_status", "ok"),
                    }
                )

            pages.append(
                {
                    "page_num": page_index + 1,
                    "page_width": page_width,
                    "page_height": page_height,
                    "blocks": blocks,
                }
            )

        return pages
    finally:
        output_document.close()
        if source_document is not None:
            source_document.close()


def _append_docx_preview_page(
    pages: list[dict[str, str]],
    page_index: int,
    source_parts: list[str],
    translated_parts: list[str],
) -> None:
    pages.append(
        {
            "id": f"section-{page_index}",
            "label": f"Section {page_index}",
            "source_text": "\n\n".join(part for part in source_parts if part).strip(),
            "translated_text": "\n\n".join(part for part in translated_parts if part).strip(),
        }
    )


def _build_docx_pages(source_path: Path, output_path: Path) -> list[dict[str, str]]:
    source_paragraphs = _docx_paragraph_texts(source_path)
    translated_paragraphs = _docx_paragraph_texts(output_path)
    total = max(len(source_paragraphs), len(translated_paragraphs))
    if total == 0:
        return [
            {
                "id": "section-1",
                "label": "Section 1",
                "source_text": "",
                "translated_text": "",
            }
        ]

    pages: list[dict[str, str]] = []
    source_parts: list[str] = []
    translated_parts: list[str] = []
    char_count = 0

    for index in range(total):
        source_text = source_paragraphs[index] if index < len(source_paragraphs) else ""
        translated_text = translated_paragraphs[index] if index < len(translated_paragraphs) else ""
        source_parts.append(source_text)
        translated_parts.append(translated_text)
        char_count += len(source_text) + len(translated_text)

        reached_limit = len(source_parts) >= DOCX_PREVIEW_PARAGRAPH_LIMIT or char_count >= DOCX_PREVIEW_CHAR_LIMIT
        is_last = index == total - 1
        if reached_limit or is_last:
            _append_docx_preview_page(pages, len(pages) + 1, source_parts, translated_parts)
            source_parts = []
            translated_parts = []
            char_count = 0

    return pages


def build_job_preview(job: TranslationJob, *, created_at: str | None = None) -> dict:
    if job.output_file is None:
        raise ValueError("Preview is unavailable until translation output is created")

    output_path = Path(job.output_file.storage_path)
    input_path = Path(job.input_file.storage_path)
    extension = output_path.suffix.lower()
    if extension == ".pdf":
        pages = _build_pdf_pages(input_path, output_path)
        document_kind = "pdf"
    elif extension == ".docx":
        pages = _build_docx_pages(input_path, output_path)
        document_kind = "docx"
    else:
        raise ValueError(f"Unsupported preview format: {extension}")

    now = _utcnow_iso()
    payload = {
        "schema_version": PREVIEW_SCHEMA_VERSION,
        "job_id": job.id,
        "title": job.input_file.original_name,
        "output_name": job.output_file.original_name,
        "document_kind": document_kind,
        "source_language": job.source_language,
        "target_language": job.target_language,
        "created_at": created_at or now,
        "updated_at": now,
        "pages": pages,
    }
    if document_kind == "pdf":
        payload["pdf_text_granularity"] = PDF_PREVIEW_TEXT_GRANULARITY
    return payload


def _write_preview(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with NamedTemporaryFile("w", encoding="utf-8", dir=path.parent, delete=False) as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)
        temp_path = Path(handle.name)
    temp_path.replace(path)


def _preview_matches_schema(payload: dict, extension: str) -> bool:
    schema_version = payload.get("schema_version")
    if schema_version is not None and not isinstance(schema_version, int):
        return False
    if extension == ".pdf":
        return payload.get("document_kind") == "pdf" and all(
            isinstance(page, dict)
            and "page_num" in page
            and "page_width" in page
            and "page_height" in page
            and "blocks" in page
            and all(isinstance(block, dict) and "type" in block for block in page.get("blocks", []))
            for page in payload.get("pages", [])
        )
    if extension == ".docx":
        return payload.get("document_kind") == "docx" and all(
            isinstance(page, dict)
            and "id" in page
            and "label" in page
            and "source_text" in page
            and "translated_text" in page
            for page in payload.get("pages", [])
        )
    return False


def _preview_uses_current_pdf_text_granularity(payload: dict) -> bool:
    return payload.get("document_kind") != "pdf" or (
        payload.get("schema_version") == PREVIEW_SCHEMA_VERSION
        and payload.get("pdf_text_granularity") == PDF_PREVIEW_TEXT_GRANULARITY
    )


def _preview_needs_pdf_font_metadata(payload: dict) -> bool:
    if payload.get("document_kind") != "pdf":
        return False

    for page in payload.get("pages", []):
        if not isinstance(page, dict):
            continue
        for block in page.get("blocks", []):
            if not isinstance(block, dict):
                continue
            if block.get("type") == "table":
                for cell in block.get("cells", []):
                    if isinstance(cell, dict) and not str(cell.get("font_name", "")).strip():
                        return True
                continue
            if not str(block.get("font_name", "")).strip():
                return True
    return False


def _cleanup_pdf_preview_text_content(preview: dict) -> bool:
    if preview.get("document_kind") != "pdf":
        return False

    changed = False
    for page in preview.get("pages", []):
        if not isinstance(page, dict):
            continue
        for block in page.get("blocks", []):
            if not isinstance(block, dict):
                continue
            if block.get("type") == "table":
                for cell in block.get("cells", []):
                    if not isinstance(cell, dict):
                        continue
                    original_size = float(cell.get("font_size_original") or 0.0)
                    current_size = float(cell.get("font_size_current") or 0.0)
                    if original_size > current_size:
                        cell["font_size_current"] = original_size
                        changed = True
                    for field_name in ("src_text", "tgt_text"):
                        sanitized_text, field_changed = _sanitize_pdf_preview_edit_text(str(cell.get(field_name, "")))
                        if field_changed:
                            cell[field_name] = sanitized_text
                            changed = True
                    continue

            original_size = float(block.get("font_size_original") or 0.0)
            current_size = float(block.get("font_size_current") or 0.0)
            if original_size > current_size:
                block["font_size_current"] = original_size
                changed = True
            for field_name in ("src_text", "tgt_text"):
                sanitized_text, field_changed = _sanitize_pdf_preview_edit_text(str(block.get(field_name, "")))
                if field_changed:
                    block[field_name] = sanitized_text
                    changed = True

    if changed:
        preview["schema_version"] = PREVIEW_SCHEMA_VERSION
        preview["updated_at"] = _utcnow_iso()
    return changed


def _enrich_pdf_preview_font_metadata(preview: dict, job: TranslationJob) -> bool:
    if job.output_file is None:
        raise ValueError("Preview is unavailable until translation output is created")

    output_path = Path(job.output_file.storage_path)
    if not output_path.exists():
        return False

    source_path = Path(job.input_file.storage_path)
    source_document = fitz.open(source_path) if source_path.exists() else None
    output_document = fitz.open(output_path)
    changed = False

    try:
        for page_index, preview_page in enumerate(preview.get("pages", [])):
            if not isinstance(preview_page, dict):
                continue

            output_page = output_document.load_page(page_index) if page_index < output_document.page_count else None
            source_page = (
                source_document.load_page(page_index)
                if source_document is not None and page_index < source_document.page_count
                else None
            )
            font_fragments = _extract_pdf_text_fragments(output_page) if output_page is not None else []
            if not font_fragments and source_page is not None:
                font_fragments = _extract_pdf_text_fragments(source_page)

            for block in preview_page.get("blocks", []):
                if not isinstance(block, dict):
                    continue

                if block.get("type") == "table":
                    for cell in block.get("cells", []):
                        if not isinstance(cell, dict) or str(cell.get("font_name", "")).strip():
                            continue
                        rect = fitz.Rect(cell.get("rect") or (0, 0, 0, 0))
                        cell["font_name"] = _dominant_font_name_in_rect(
                            font_fragments,
                            rect,
                            fallback_text=str(cell.get("tgt_text") or cell.get("src_text") or ""),
                        )
                        changed = True
                    continue

                if str(block.get("font_name", "")).strip():
                    continue
                rect = fitz.Rect(block.get("rect") or (0, 0, 0, 0))
                block["font_name"] = _dominant_font_name_in_rect(
                    font_fragments,
                    rect,
                    fallback_text=str(block.get("tgt_text") or block.get("src_text") or ""),
                )
                changed = True
    finally:
        output_document.close()
        if source_document is not None:
            source_document.close()

    if changed:
        preview["schema_version"] = PREVIEW_SCHEMA_VERSION
        preview["updated_at"] = _utcnow_iso()
    return changed


def _copy_preview_table_blocks(page: dict[str, object]) -> list[dict[str, object]]:
    return [deepcopy(block) for block in page.get("blocks", []) if isinstance(block, dict) and block.get("type") == "table"]


def _build_pdf_pages_from_existing_preview(source_path: Path, output_path: Path, existing_pages: list[dict[str, object]]) -> list[dict[str, object]]:
    source_document = fitz.open(source_path) if source_path.exists() else None
    output_document = fitz.open(output_path)
    try:
        page_count = max(source_document.page_count if source_document is not None else 0, output_document.page_count, len(existing_pages), 1)
        pages: list[dict[str, object]] = []

        for page_index in range(page_count):
            existing_page = existing_pages[page_index] if page_index < len(existing_pages) else {}
            source_page = (
                source_document.load_page(page_index)
                if source_document is not None and page_index < source_document.page_count
                else None
            )
            output_page = output_document.load_page(page_index) if page_index < output_document.page_count else None
            geometry_page = output_page or source_page
            page_width = round(geometry_page.rect.width, 2) if geometry_page is not None else float(existing_page.get("page_width", 0.0) or 0.0)
            page_height = round(geometry_page.rect.height, 2) if geometry_page is not None else float(existing_page.get("page_height", 0.0) or 0.0)
            source_fragments = _extract_pdf_text_fragments(source_page) if source_page is not None else []
            target_fragments = _extract_pdf_text_fragments(output_page) if output_page is not None else []

            table_blocks = _copy_preview_table_blocks(existing_page if isinstance(existing_page, dict) else {})
            table_rects = [fitz.Rect(block["table_rect"]) for block in table_blocks if isinstance(block.get("table_rect"), list)]
            filtered_source_fragments = [fragment for fragment in source_fragments if not _fragment_belongs_to_table(fragment, table_rects)]
            filtered_target_fragments = [fragment for fragment in target_fragments if not _fragment_belongs_to_table(fragment, table_rects)]
            geometry_fragments = filtered_source_fragments or filtered_target_fragments

            text_blocks = _build_pdf_text_blocks(
                geometry_fragments,
                page_index=page_index,
                source_fragments=filtered_source_fragments,
                target_fragments=filtered_target_fragments,
                page_width=page_width,
            )
            page_items = sorted([*text_blocks, *table_blocks], key=_page_item_sort_key)

            blocks: list[dict[str, object]] = []
            text_block_index = 0
            for block in page_items:
                if block.get("type") == "table":
                    blocks.append(block)
                    continue

                text_block_index += 1
                blocks.append(
                    {
                        "type": "text",
                        "block_id": f"p{page_index + 1}_b{text_block_index}",
                        "rect": block["rect"],
                        "font_name": str(block.get("font_name", "")),
                        "font_size_original": float(block.get("font_size_original", 12.0) or 12.0),
                        "font_size_current": float(block.get("font_size_current", 12.0) or 12.0),
                        "src_text": str(block.get("src_text", "")),
                        "tgt_text": str(block.get("tgt_text", "")),
                        "alignment": block.get("alignment"),
                        "font_style": block.get("font_style"),
                        "rotation": block.get("rotation"),
                        "layout_status": block.get("layout_status", "ok"),
                    }
                )

            pages.append(
                {
                    "page_num": page_index + 1,
                    "page_width": page_width,
                    "page_height": page_height,
                    "blocks": blocks,
                }
            )

        return pages
    finally:
        output_document.close()
        if source_document is not None:
            source_document.close()


def _migrate_pdf_preview_payload(preview: dict, job: TranslationJob) -> dict:
    if job.output_file is None:
        raise ValueError("Preview is unavailable until translation output is created")

    migrated_preview = dict(preview)
    migrated_preview["schema_version"] = PREVIEW_SCHEMA_VERSION
    migrated_preview["pdf_text_granularity"] = PDF_PREVIEW_TEXT_GRANULARITY
    migrated_preview["updated_at"] = _utcnow_iso()
    migrated_preview["pages"] = _build_pdf_pages_from_existing_preview(
        Path(job.input_file.storage_path),
        Path(job.output_file.storage_path),
        [page for page in preview.get("pages", []) if isinstance(page, dict)],
    )
    return migrated_preview


def load_or_create_preview(job: TranslationJob, *, force: bool = False, migrate_pdf_text_blocks: bool = True) -> dict:
    if job.output_file is None:
        raise ValueError("Preview is unavailable until translation output is created")

    output_path = Path(job.output_file.storage_path)
    sidecar_path = preview_sidecar_path(job.output_file.storage_path)
    existing_preview: dict | None = None
    if sidecar_path.exists():
        try:
            with sidecar_path.open("r", encoding="utf-8") as handle:
                loaded = json.load(handle)
        except json.JSONDecodeError:
            loaded = None
        if isinstance(loaded, dict):
            existing_preview = loaded

    extension = output_path.suffix.lower()
    if not force and existing_preview is not None and _preview_matches_schema(existing_preview, extension):
        preview_changed = False
        if migrate_pdf_text_blocks and extension == ".pdf" and not _preview_uses_current_pdf_text_granularity(existing_preview):
            existing_preview = _migrate_pdf_preview_payload(existing_preview, job)
            preview_changed = True
        if extension == ".pdf":
            preview_changed = _cleanup_pdf_preview_text_content(existing_preview) or preview_changed
        if extension == ".pdf" and _preview_needs_pdf_font_metadata(existing_preview):
            preview_changed = _enrich_pdf_preview_font_metadata(existing_preview, job) or preview_changed
        if preview_changed:
            _write_preview(sidecar_path, existing_preview)
        return existing_preview

    payload = build_job_preview(job, created_at=existing_preview.get("created_at") if existing_preview else None)
    _write_preview(sidecar_path, payload)
    return payload


def _update_docx_preview(preview: dict, page_updates: list[dict[str, str]]) -> dict:
    page_lookup = {page["id"]: page for page in preview["pages"]}
    for update in page_updates:
        page_id = update["id"]
        if page_id not in page_lookup:
            raise ValueError(f"Preview page '{page_id}' does not exist")
        page_lookup[page_id]["translated_text"] = update["translated_text"]

    preview["updated_at"] = _utcnow_iso()
    return preview


def _ensure_pdf_render_font(page: fitz.Page, font: PdfRenderFont) -> None:
    if font.render_buffer:
        page.insert_font(fontname=font.render_name, fontbuffer=font.render_buffer)
        return
    if font.render_file:
        page.insert_font(fontname=font.render_name, fontfile=font.render_file)


def _insert_pdf_text(
    page: fitz.Page,
    point: tuple[float, float],
    text: str,
    font_size: float,
    font: PdfRenderFont,
    *,
    rotate: int = 0,
) -> None:
    x, y = point
    for run_text, run_font in _pdf_text_font_runs(text, font):
        _ensure_pdf_render_font(page, run_font)
        insert_kwargs = {
            "fontsize": font_size,
            "fontname": run_font.render_name,
            "overlay": True,
        }
        if run_font.render_file:
            insert_kwargs["fontfile"] = run_font.render_file
        if rotate:
            insert_kwargs["rotate"] = rotate
        page.insert_text((x, y), run_text, **insert_kwargs)
        x += _pdf_text_width(run_text, run_font, font_size)


def _insert_pdf_text_single_font(
    page: fitz.Page,
    point: tuple[float, float],
    text: str,
    font_size: float,
    font: PdfRenderFont,
) -> None:
    _ensure_pdf_render_font(page, font)
    insert_kwargs = {
        "fontsize": font_size,
        "fontname": font.render_name,
        "overlay": True,
    }
    if font.render_file:
        insert_kwargs["fontfile"] = font.render_file
    page.insert_text(point, text, **insert_kwargs)


def _insert_pdf_textbox(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    font_size: float,
    font: PdfRenderFont,
    *,
    align: int = 0,
) -> float:
    runs = _pdf_text_font_runs(text, font)
    if len(runs) > 1 or runs[0][1] != font:
        return -1.0
    insert_kwargs = {
        "fontsize": font_size,
        "fontname": font.render_name,
        "overlay": True,
        "align": align,
    }
    if font.render_file:
        insert_kwargs["fontfile"] = font.render_file
    return float(page.insert_textbox(rect, text, **insert_kwargs))


def _insert_pdf_textbox_single_font(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    font_size: float,
    font: PdfRenderFont,
    *,
    align: int = 0,
) -> float:
    _ensure_pdf_render_font(page, font)
    insert_kwargs = {
        "fontsize": font_size,
        "fontname": font.render_name,
        "overlay": True,
        "align": align,
    }
    if font.render_file:
        insert_kwargs["fontfile"] = font.render_file
    return float(page.insert_textbox(rect, text, **insert_kwargs))


def _force_insert_pdf_text_lines(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    font_size: float,
    font: PdfRenderFont,
    *,
    max_lines: int = 1,
    min_font_size: float = PDF_MIN_REDRAW_FONT_SIZE,
    step_down: float = 0.5,
    line_height_multiplier: float = 1.15,
    align: int = 0,
    rotate: int = 0,
) -> float:
    line_limit = max(1, int(max_lines or 1))
    minimum_size = max(float(min_font_size or PDF_MIN_REDRAW_FONT_SIZE), PDF_MIN_REDRAW_FONT_SIZE)
    current_size = round(max(float(font_size or minimum_size), minimum_size), 2)
    while current_size >= minimum_size:
        line_step = max(current_size * line_height_multiplier, current_size + 1.0)
        line_count = min(line_limit, max(1, int(max(rect.height, line_step) // line_step)))
        wrapped_lines = _wrap_pdf_text_to_widths(text, [max(rect.width - 1.0, 1.0)] * line_count, font, current_size)
        if wrapped_lines is not None:
            for index, line_text in enumerate(wrapped_lines):
                if not line_text:
                    continue
                text_width = _pdf_text_width(line_text, font, current_size)
                baseline_x = rect.x0
                if align == 1:
                    baseline_x = max(rect.x0, rect.x0 + (rect.width - text_width) / 2)
                baseline_y = min(rect.y0 + current_size + index * line_step, rect.y1 - 0.5)
                _insert_pdf_text(page, (baseline_x, baseline_y), line_text, current_size, font, rotate=rotate)
            return current_size
        current_size = round(current_size - step_down, 2)

    _insert_pdf_text(page, (rect.x0, max(rect.y0 + minimum_size, rect.y1 - 0.5)), text, minimum_size, font, rotate=rotate)
    return minimum_size


def _draw_pdf_text_into_line_rects(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    font_size: float,
    line_fragments: list[PdfPreviewFragment],
    render_font: PdfRenderFont,
    fallback_fragments: list[PdfPreviewFragment] | None = None,
    clear_full_rect: bool = False,
    min_font_size: float = PDF_MIN_REDRAW_FONT_SIZE,
    step_down: float = 0.5,
    line_height_multiplier: float = 1.12,
    align: int = 0,
) -> tuple[str, float] | None:
    minimum_size = max(float(min_font_size or PDF_MIN_REDRAW_FONT_SIZE), PDF_MIN_REDRAW_FONT_SIZE)
    normalized_text = _normalize_pdf_block_text(text)
    if not normalized_text:
        return "", round(max(float(font_size or PDF_EDITOR_MIN_FONT_SIZE), minimum_size), 2)

    requested_size = max(float(font_size or PDF_EDITOR_MIN_FONT_SIZE), minimum_size)
    for full_width in (False, True):
        current_size = round(requested_size, 2)
        while current_size >= minimum_size:
            line_rects, base_count = _build_pdf_line_rects(
                rect,
                line_fragments,
                current_size,
                full_width=full_width,
                line_height_multiplier=line_height_multiplier,
            )
            if line_rects:
                widths = [max(line_rect.width, 1.0) for line_rect in line_rects]
                wrapped_lines = _wrap_pdf_text_to_widths(normalized_text, widths, render_font, current_size)
                if wrapped_lines is not None and len(wrapped_lines) <= len(line_rects):
                    if clear_full_rect:
                        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                    else:
                        clear_count = max(base_count, len(wrapped_lines))
                        for clear_rect in line_rects[:clear_count]:
                            page.draw_rect(clear_rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)

                    for line_rect, line_text in zip(line_rects, wrapped_lines, strict=False):
                        if not line_text:
                            continue
                        text_width = _pdf_text_width(line_text, render_font, current_size)
                        baseline_x = line_rect.x0
                        if align == 1:
                            baseline_x = max(line_rect.x0, line_rect.x0 + (line_rect.width - text_width) / 2)
                        baseline_y = min(line_rect.y1 - 0.5, line_rect.y0 + current_size)
                        _insert_pdf_text(page, (baseline_x, baseline_y), line_text, current_size, render_font)
                    return normalized_text, current_size
            current_size = round(current_size - step_down, 2)

    if fallback_fragments:
        current_size = round(requested_size, 2)
        while current_size >= minimum_size:
            line_rects, base_count = _build_pdf_line_rects(
                rect,
                fallback_fragments,
                current_size,
                full_width=True,
                line_height_multiplier=line_height_multiplier,
            )
            if line_rects:
                widths = [max(line_rect.width, 1.0) for line_rect in line_rects]
                wrapped_lines = _wrap_pdf_text_to_widths(normalized_text, widths, render_font, current_size)
                if wrapped_lines is not None and len(wrapped_lines) <= len(line_rects):
                    if clear_full_rect:
                        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                    else:
                        clear_count = max(base_count, len(wrapped_lines))
                        for clear_rect in line_rects[:clear_count]:
                            page.draw_rect(clear_rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)

                    for line_rect, line_text in zip(line_rects, wrapped_lines, strict=False):
                        if not line_text:
                            continue
                        text_width = _pdf_text_width(line_text, render_font, current_size)
                        baseline_x = line_rect.x0
                        if align == 1:
                            baseline_x = max(line_rect.x0, line_rect.x0 + (line_rect.width - text_width) / 2)
                        baseline_y = min(line_rect.y1 - 0.5, line_rect.y0 + current_size)
                        _insert_pdf_text(page, (baseline_x, baseline_y), line_text, current_size, render_font)
                    return normalized_text, current_size
            current_size = round(current_size - step_down, 2)

    return None


def _draw_pdf_block_text(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    font_size: float,
    *,
    preferred_font_name: str = "",
    page_fonts: tuple[PdfPageFontResource, ...] = (),
    prefer_external_cjk_font: bool = False,
    language_profile: PdfLanguageProfile | None = None,
    line_fragments: list[PdfPreviewFragment] | None = None,
    fallback_fragments: list[PdfPreviewFragment] | None = None,
    clear_full_rect: bool = False,
    align: int = 0,
    force_bold: bool = False,
    rotate: int = 0,
) -> tuple[str, float]:
    minimum_size = language_profile.min_font_size if language_profile is not None else PDF_MIN_REDRAW_FONT_SIZE
    step_down = language_profile.step_down if language_profile is not None else 0.5
    line_height_multiplier = language_profile.line_height_multiplier if language_profile is not None else 1.12
    normalized_text = _normalize_pdf_block_text(text)
    requested_size = max(float(font_size or minimum_size), minimum_size)
    if not normalized_text:
        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
        return "", round(requested_size, 2)

    render_font = _resolve_pdf_render_font(
        normalized_text,
        preferred_font_name=preferred_font_name,
        page_fonts=page_fonts,
        prefer_external_cjk_font=prefer_external_cjk_font,
        language_profile=language_profile,
        force_bold=force_bold,
    )
    _ensure_pdf_render_font(page, render_font)
    if line_fragments and rotate not in {90, 270}:
        line_drawn = _draw_pdf_text_into_line_rects(
            page,
            rect,
            normalized_text,
            requested_size,
            line_fragments,
            render_font,
            fallback_fragments=fallback_fragments,
            clear_full_rect=clear_full_rect,
            min_font_size=minimum_size,
            step_down=step_down,
            line_height_multiplier=line_height_multiplier,
            align=align,
        )
        if line_drawn is not None:
            return line_drawn

    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
    if rotate in {90, 270}:
        applied_font_size = _force_insert_pdf_text_lines(
            page,
            rect,
            normalized_text,
            requested_size,
            render_font,
            max_lines=len(line_fragments or fallback_fragments or ()) or 1,
            min_font_size=minimum_size,
            step_down=step_down,
            line_height_multiplier=line_height_multiplier,
            align=align,
            rotate=rotate,
        )
        return normalized_text, applied_font_size

    current_size = round(requested_size, 2)
    while current_size >= minimum_size:
        remainder = _insert_pdf_textbox(page, rect, normalized_text, current_size, render_font, align=align)
        if remainder >= 0:
            return normalized_text, current_size
        current_size = round(current_size - step_down, 2)

    fallback_line_count = len(line_fragments or fallback_fragments or ()) or 1
    applied_font_size = _force_insert_pdf_text_lines(
        page,
        rect,
        normalized_text,
        minimum_size,
        render_font,
        max_lines=fallback_line_count,
        min_font_size=minimum_size,
        step_down=step_down,
        line_height_multiplier=line_height_multiplier,
        align=align,
    )
    return normalized_text, applied_font_size


def _apply_pdf_preview_updates(job: TranslationJob, preview: dict, block_updates: list[dict[str, object]]) -> dict[str, dict[str, object]]:
    if job.output_file is None:
        raise ValueError("Translated PDF does not exist")
    for update in block_updates:
        if str(update.get("layout_status") or "").casefold() == "overflow":
            target_id = str(update.get("cell_id") or update.get("block_id") or "")
            raise ValueError(f"Preview block '{target_id}' still overflows. Resolve it before saving.")

    output_path = Path(job.output_file.storage_path)
    if not output_path.exists():
        raise ValueError("Translated PDF does not exist")

    editable_lookup: dict[str, tuple[int, dict[str, object], str]] = {}
    table_font_floor_by_cell_id: dict[str, float] = {}
    applied_updates: dict[str, dict[str, object]] = {}
    language_profile = _pdf_language_profile(job.target_language)
    for page in preview.get("pages", []):
        page_index = int(page["page_num"]) - 1
        for block in page.get("blocks", []):
            if block.get("type") == "table":
                table_font_floor = minimum_sibling_font_size(block.get("cells", []), PDF_MIN_REDRAW_FONT_SIZE)
                for cell in block.get("cells", []):
                    editable_lookup[str(cell["cell_id"])] = (page_index, cell, "cell")
                    table_font_floor_by_cell_id[str(cell["cell_id"])] = table_font_floor
                continue
            editable_lookup[str(block["block_id"])] = (page_index, block, "block")

    document = fitz.open(output_path)
    source_path = Path(job.input_file.storage_path)
    source_document = fitz.open(source_path) if source_path.exists() else None
    with NamedTemporaryFile("wb", suffix=output_path.suffix, dir=output_path.parent, delete=False) as handle:
        temp_path = Path(handle.name)

    try:
        page_cache: dict[int, fitz.Page] = {}
        page_font_cache: dict[int, tuple[PdfPageFontResource, ...]] = {}
        page_fragment_cache: dict[int, list[PdfPreviewFragment]] = {}
        source_fragment_cache: dict[int, list[PdfPreviewFragment]] = {}
        language_route = font_route_for_language(job.target_language)

        for update in block_updates:
            target_id = str(update.get("cell_id") or update.get("block_id") or "")
            lookup = editable_lookup.get(target_id)
            if lookup is None:
                raise ValueError(f"Preview block '{target_id}' does not exist")

            page_index, block, block_kind = lookup
            page = page_cache.get(page_index)
            if page is None:
                page = document.load_page(page_index)
                page_cache[page_index] = page
                page_font_cache[page_index] = _extract_page_font_resources(page)
                page_fragment_cache[page_index] = _extract_pdf_text_fragments(page)
            if page_index not in source_fragment_cache:
                source_page = (
                    source_document.load_page(page_index)
                    if source_document is not None and page_index < source_document.page_count
                    else None
                )
                source_fragment_cache[page_index] = _extract_pdf_text_fragments(source_page) if source_page is not None else []
            rect = fitz.Rect(block["rect"])
            page_fonts = page_font_cache.get(page_index, ())
            page_fragments = page_fragment_cache.get(page_index, [])
            source_fragments = source_fragment_cache.get(page_index, [])
            preferred_font_name = str(
                block.get("font_name")
                or _dominant_font_name_in_rect(
                    page_fragments,
                    rect,
                    fallback_text=str(update.get("tgt_text") or block.get("tgt_text") or block.get("src_text") or ""),
                )
            )
            source_line_fragments = _line_fragments_in_rect(source_fragments, rect)
            page_line_fragments = _line_fragments_in_rect(page_fragments, rect)
            line_fragments = source_line_fragments or page_line_fragments or None
            fallback_fragments = page_line_fragments or source_line_fragments or None
            prefer_external_cjk_font = True
            align = 1 if block.get("alignment") == "CENTER" else 0
            force_bold = block.get("font_style") == "BOLD"
            rotation = int(block.get("rotation") or 0)
            requested_font_size = max(
                float(update.get("font_size_final") or 0.0),
                float(block.get("font_size_original") or 0.0),
                float(block.get("font_size_current") or 0.0),
                PDF_MIN_REDRAW_FONT_SIZE,
            )
            if block_kind == "cell":
                requested_font_size = min(requested_font_size, table_font_floor_by_cell_id.get(target_id, requested_font_size))
            if language_route is not None:
                requested_font_size = max(PDF_MIN_REDRAW_FONT_SIZE, requested_font_size * language_route.scale)
            visible_text = _normalize_pdf_block_text(str(update.get("tgt_text", "")))
            shielded_text = apply_thai_word_wrap_shield(str(update.get("tgt_text", "")), job.target_language)
            normalized_text, applied_font_size = _draw_pdf_block_text(
                page,
                rect,
                shielded_text,
                requested_font_size,
                preferred_font_name=preferred_font_name,
                page_fonts=page_fonts,
                prefer_external_cjk_font=prefer_external_cjk_font,
                language_profile=language_profile,
                line_fragments=line_fragments,
                fallback_fragments=fallback_fragments,
                clear_full_rect=block_kind != "cell",
                align=align,
                force_bold=force_bold,
                rotate=rotation,
            )
            applied_updates[target_id] = {
                "tgt_text": visible_text,
                "font_size_current": applied_font_size,
                "font_name": preferred_font_name,
                "layout_status": str(update.get("layout_status") or block.get("layout_status") or "ok"),
            }

        document.subset_fonts(fallback=True)
        document.save(temp_path, garbage=4, deflate=True, clean=True, deflate_fonts=True)
    except Exception:
        if temp_path.exists():
            temp_path.unlink()
        raise
    finally:
        document.close()
        if source_document is not None:
            source_document.close()

    temp_path.replace(output_path)
    job.output_file.size_bytes = output_path.stat().st_size
    job.output_file.checksum = file_checksum(output_path)
    return applied_updates


def _update_pdf_preview(preview: dict, applied_updates: dict[str, dict[str, object]]) -> dict:
    editable_lookup: dict[str, dict[str, object]] = {}
    for page in preview.get("pages", []):
        for block in page.get("blocks", []):
            if block.get("type") == "table":
                for cell in block.get("cells", []):
                    editable_lookup[str(cell["cell_id"])] = cell
                continue
            editable_lookup[str(block["block_id"])] = block

    for target_id, applied in applied_updates.items():
        editable = editable_lookup.get(target_id)
        if editable is None:
            raise ValueError(f"Preview block '{target_id}' does not exist")
        editable["tgt_text"] = str(applied.get("tgt_text", ""))
        editable["font_size_current"] = float(applied.get("font_size_current") or editable.get("font_size_current") or 12.0)
        editable["layout_status"] = str(applied.get("layout_status") or editable.get("layout_status") or "ok")
        if applied.get("font_name"):
            editable["font_name"] = str(applied["font_name"])

    preview["updated_at"] = _utcnow_iso()
    return preview


def update_preview(job: TranslationJob, update_payload: dict) -> dict:
    preview = load_or_create_preview(job, migrate_pdf_text_blocks=False)
    sidecar = preview_sidecar_path(job.output_file.storage_path)

    if preview["document_kind"] == "pdf":
        if update_payload.get("status") != "validated":
            raise ValueError("PDF preview updates require status 'validated'")
        block_updates = update_payload.get("payload")
        if not isinstance(block_updates, list):
            raise ValueError("PDF preview updates require a block payload")
        if not _preview_uses_current_pdf_text_granularity(preview):
            migrated_preview = _migrate_pdf_preview_payload(preview, job)
            _write_preview(sidecar, migrated_preview)
            raise ValueError("Preview layout was upgraded. Please reopen the preview and save again.")

        applied_updates = _apply_pdf_preview_updates(job, preview, block_updates)
        updated_preview = _update_pdf_preview(preview, applied_updates)
        _write_preview(sidecar, updated_preview)
        return updated_preview

    page_updates = update_payload.get("pages")
    if not isinstance(page_updates, list):
        raise ValueError("DOCX preview updates require page content")

    updated_preview = _update_docx_preview(preview, page_updates)
    _write_preview(sidecar, updated_preview)
    return updated_preview
