from __future__ import annotations

import json
import logging
import os
import re
import shutil
import tempfile
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

import fitz
import httpx
import pytesseract
from docx import Document
from PIL import Image
from sqlalchemy.orm import Session, selectinload

from doc_translator.audit import record_audit
from doc_translator.babeldoc_hooks import babeldoc_ir_sidecar_path
from doc_translator.babeldoc_runner import BabeldocLibraryResult, translate_pdf_with_babeldoc_library
from doc_translator.db import SessionLocal
from doc_translator.models import JobEvent, JobFile, JobFileKind, JobStatus, TranslationJob
from doc_translator.preview import load_or_create_preview
from doc_translator.settings_service import RuntimeSettings, get_runtime_settings
from doc_translator.storage import build_output_target, file_checksum
from doc_translator.translators.prompt_builder import build_terminology_instruction
from doc_translator.translators.trie_matcher import default_terminology_matcher


logger = logging.getLogger(__name__)

BABELDOC_LANGUAGE_CODES = {
    "auto": None,
    "auto detect": None,
    "en": "English",
    "english": "English",
    "zh": "zh-CN",
    "zh-cn": "zh-CN",
    "chinese": "zh-CN",
    "ja": "Japanese",
    "japanese": "Japanese",
    "ko": "Korean",
    "korean": "Korean",
    "ms": "Malay",
    "malay": "Malay",
    "th": "Thai",
    "thai": "Thai",
    "vi": "Vietnamese",
    "vietnamese": "Vietnamese",
    "es": "es",
    "spanish": "es",
    "fr": "fr",
    "french": "fr",
    "de": "de",
    "german": "de",
}
PDF_OCR_DPI = 300
PDF_BACKGROUND_ANALYSIS_DPI = 24
PDF_OCR_WORKAROUND_MIN_LUMINANCE = 235.0
PDF_OCR_WORKAROUND_MIN_BRIGHT_RATIO = 0.6
PDF_OCR_WORKAROUND_PAGE_SAMPLE_LIMIT = 5
PDF_OCR_WORKAROUND_MAX_AVERAGE_DRAWING_ITEMS = 80.0
PDF_OCR_WORKAROUND_MAX_AVERAGE_LINE_DRAWINGS = 25.0
PDF_OCR_TEXT_MASK_HORIZONTAL_PADDING = 1.5
PDF_OCR_TEXT_MASK_VERTICAL_PADDING = 1.0
BABELDOC_QPS = max(1, int(os.getenv("BABELDOC_QPS", "6")))
BABELDOC_PROGRESS_REPORT_INTERVAL_SECONDS = 0.5
_BABELDOC_PROGRESS_START = 20
_BABELDOC_PROGRESS_END = 92
_BABELDOC_PROGRESS_EVENT_PREFIX = "__BABELDOC_PROGRESS__"
_ANSI_ESCAPE_RE = re.compile(r"\x1b(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])")
_BABELDOC_PROGRESS_FRACTION_RE = re.compile(r"(\d+)\s*/\s*(\d+)")
_MODEL_ERROR_BODY_LIMIT = 600


@dataclass(frozen=True)
class _BabeldocProgressStage:
    name: str
    weight: float
    status: JobStatus
    message: str
    cumulative_weight_before: float


_BABELDOC_PROGRESS_STAGE_BLUEPRINTS = [
    ("Parse PDF and Create Intermediate Representation", 14.12, JobStatus.PARSING, "Analyzing PDF structure"),
    ("DetectScannedFile", 2.45, JobStatus.PARSING, "Checking whether PDF needs OCR"),
    ("Parse Page Layout", 14.03, JobStatus.PARSING, "Analyzing page layout"),
    ("Parse Table", 1.0, JobStatus.PARSING, "Analyzing tables"),
    ("Parse Paragraphs", 6.26, JobStatus.PARSING, "Analyzing paragraphs"),
    ("Parse Formulas and Styles", 1.66, JobStatus.PARSING, "Analyzing formulas and styles"),
    ("Automatic Term Extraction", 30.0, JobStatus.TRANSLATING, "Extracting terminology"),
    ("Translate Paragraphs", 46.96, JobStatus.TRANSLATING, "Translating extracted text"),
    ("Typesetting", 4.71, JobStatus.REBUILDING, "Typesetting translated PDF"),
    ("Add Fonts", 0.61, JobStatus.REBUILDING, "Embedding fonts"),
    ("Generate drawing instructions", 1.96, JobStatus.REBUILDING, "Generating drawing instructions"),
    ("Subset font", 0.92, JobStatus.REBUILDING, "Optimizing embedded fonts"),
    ("Save PDF", 6.34, JobStatus.REBUILDING, "Saving translated PDF"),
]


def _build_babeldoc_progress_stages() -> tuple[_BabeldocProgressStage, ...]:
    cumulative_weight = 0.0
    stages: list[_BabeldocProgressStage] = []
    for name, weight, status, message in _BABELDOC_PROGRESS_STAGE_BLUEPRINTS:
        stages.append(
            _BabeldocProgressStage(
                name=name,
                weight=weight,
                status=status,
                message=message,
                cumulative_weight_before=cumulative_weight,
            )
        )
        cumulative_weight += weight
    return tuple(stages)


_BABELDOC_PROGRESS_STAGES = _build_babeldoc_progress_stages()
_BABELDOC_PROGRESS_STAGE_BY_NAME = {stage.name: stage for stage in _BABELDOC_PROGRESS_STAGES}
_BABELDOC_TOTAL_STAGE_WEIGHT = sum(stage.weight for stage in _BABELDOC_PROGRESS_STAGES)


class JobCancelledError(Exception):
    pass


class OpenAICompatibleTranslator:
    def __init__(self, runtime: RuntimeSettings) -> None:
        self.runtime = runtime
        base_url = runtime.model_base_url.rstrip("/")
        self.endpoint = f"{base_url}/chat/completions"

    def translate_text(
        self,
        text: str,
        source_language: str,
        target_language: str,
        *,
        preserve_line_breaks: bool = True,
        extra_system_instruction: str = "",
    ) -> str:
        if not text.strip():
            return text
        formatting_instruction = (
            "Preserve line breaks and lists."
            if preserve_line_breaks
            else "Preserve paragraph breaks and lists, but reflow ordinary line breaks naturally for the target language."
        )
        system_content = (
            f"Translate the user's text from {source_language} to {target_language}. "
            f"Return only the translated text. {formatting_instruction} "
            "Keep citations, numbers, and inline Latin-script terms accurately formatted. "
            "Do not expand or explain short technical expressions or parameter mnemonics; "
            "keep forms such as 'FS =' unchanged instead of rewriting them as explanatory prose."
        )
        if extra_system_instruction:
            system_content = f"{system_content}\n\n{extra_system_instruction}"

        response = httpx.post(
            self.endpoint,
            headers={
                "Authorization": f"Bearer {self.runtime.model_api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": self.runtime.model_name,
                "temperature": 0,
                "messages": [
                    {
                        "role": "system",
                        "content": system_content,
                    },
                    {"role": "user", "content": text},
                ],
            },
            timeout=self.runtime.model_timeout_seconds,
        )
        try:
            response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            body = response.text.strip()
            summary = re.sub(r"\s+", " ", body)[:_MODEL_ERROR_BODY_LIMIT].strip()
            reason = response.reason_phrase or "HTTP error"
            if summary:
                raise RuntimeError(
                    f"Model API request failed with {response.status_code} {reason}: {summary}"
                ) from exc
            raise RuntimeError(f"Model API request failed with {response.status_code} {reason}") from exc
        payload = response.json()
        try:
            return payload["choices"][0]["message"]["content"].strip()
        except (KeyError, IndexError, TypeError) as exc:
            raise RuntimeError("Model API returned an unexpected response") from exc

    def test_connection(self) -> tuple[int, str]:
        started = datetime.now(timezone.utc)
        preview = self.translate_text("hello", "English", "Spanish")
        latency = int((datetime.now(timezone.utc) - started).total_seconds() * 1000)
        return latency, preview[:120]


def add_job_event(session: Session, job: TranslationJob, message: str, *, level: str = "info", details: dict | None = None) -> None:
    session.add(JobEvent(job_id=job.id, level=level, message=message, details=details))


def update_job_state(
    session: Session,
    job: TranslationJob,
    *,
    status: JobStatus | None = None,
    progress: int | None = None,
    error_message: str | None = None,
    message: str | None = None,
    details: dict | None = None,
) -> None:
    if status is not None:
        job.status = status
    if progress is not None:
        job.progress = progress
    if error_message is not None:
        job.error_message = error_message
    if status == JobStatus.PARSING and job.started_at is None:
        job.started_at = datetime.now(timezone.utc)
    if status in {JobStatus.COMPLETED, JobStatus.FAILED, JobStatus.CANCELLED}:
        job.completed_at = datetime.now(timezone.utc)
    if message:
        add_job_event(session, job, message, level="error" if status == JobStatus.FAILED else "info", details=details)
    session.commit()


def ensure_not_cancelled(session: Session, job: TranslationJob) -> None:
    session.refresh(job)
    if job.cancel_requested or job.status == JobStatus.CANCELLED:
        raise JobCancelledError("Job was cancelled")


def split_text(text: str, max_chars: int = 1800) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    parts: list[str] = []
    current = ""
    for paragraph in text.splitlines(keepends=True):
        if len(current) + len(paragraph) > max_chars and current:
            parts.append(current)
            current = paragraph
        else:
            current += paragraph
    if current:
        parts.append(current)
    return parts


def translate_segments(
    translator: OpenAICompatibleTranslator,
    segments: list[str],
    *,
    source_language: str,
    target_language: str,
    preserve_line_breaks: bool,
    on_progress: Callable[[int, int], None],
    cancel_check: Callable[[], None],
) -> list[str]:
    translated: list[str] = []
    total = len(segments)
    for index, segment in enumerate(segments, start=1):
        cancel_check()
        chunked = split_text(segment)
        translated_chunks: list[str] = []
        for chunk in chunked:
            terminology_instruction = build_terminology_instruction(default_terminology_matcher().scan([chunk]))
            translated_chunks.append(
                translator.translate_text(
                    chunk,
                    source_language=source_language,
                    target_language=target_language,
                    preserve_line_breaks=preserve_line_breaks,
                    extra_system_instruction=terminology_instruction,
                )
            )
        translated.append("".join(translated_chunks))
        on_progress(index, total)
    return translated


def _pixmap_to_image(pixmap: fitz.Pixmap) -> Image.Image:
    mode = "RGBA" if pixmap.alpha else "RGB"
    return Image.frombytes(mode, [pixmap.width, pixmap.height], pixmap.samples)


def _normalize_language(language: str) -> str:
    return language.strip().lower()


def _babeldoc_language_code(language: str, *, allow_auto: bool) -> str | None:
    normalized = _normalize_language(language)
    code = BABELDOC_LANGUAGE_CODES.get(normalized)
    if normalized not in BABELDOC_LANGUAGE_CODES:
        supported = "English, Chinese, Japanese, Korean, Malay, Thai, Vietnamese, Spanish, French, German"
        raise RuntimeError(f"Unsupported PDF language '{language}'. Supported values: {supported}")
    if code is None:
        if allow_auto:
            return None
        raise RuntimeError("PDF target language must be explicitly selected")
    return code


def _page_has_extractable_text(page: fitz.Page) -> bool:
    return bool(page.get_text("words"))


def _page_has_text_resources(page: fitz.Page) -> bool:
    return bool(page.get_fonts(full=True))


def _ocr_language(runtime: RuntimeSettings) -> str | None:
    return None if runtime.ocr_language_hint == "auto" else runtime.ocr_language_hint


def _has_codepoint_in_ranges(text: str, ranges: tuple[tuple[str, str], ...]) -> bool:
    return any(start <= char <= end for char in text for start, end in ranges)


def _select_pdf_font(text: str) -> str:
    if _has_codepoint_in_ranges(text, (("\uac00", "\ud7af"), ("\u1100", "\u11ff"), ("\u3130", "\u318f"))):
        return "korea"
    if _has_codepoint_in_ranges(text, (("\u3040", "\u30ff"), ("\u31f0", "\u31ff"))):
        return "japan"
    if _has_codepoint_in_ranges(
        text,
        (
            ("\u3400", "\u4dbf"),
            ("\u4e00", "\u9fff"),
            ("\uf900", "\ufaff"),
            ("\uff00", "\uffef"),
            ("\u0400", "\u04ff"),
        ),
    ):
        return "china-s"
    return "helv"


def _insert_background_text_best_fit(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
) -> None:
    if not text.strip():
        return

    font_name = _select_pdf_font(text)
    initial_font_size = max(6.0, min(rect.height * 0.9, 22.0))
    font_sizes = [initial_font_size, max(initial_font_size - 1, 6.0), max(initial_font_size - 2, 6.0), 6.0]
    tried_sizes: set[float] = set()
    for font_size in font_sizes:
        rounded_size = round(font_size, 2)
        if rounded_size in tried_sizes:
            continue
        tried_sizes.add(rounded_size)
        remainder = page.insert_textbox(
            rect,
            text,
            fontsize=rounded_size,
            fontname=font_name,
            overlay=False,
        )
        if remainder >= 0:
            return

    page.insert_text(
        rect.bottom_left,
        text,
        fontsize=6,
        fontname=font_name,
        overlay=False,
    )


def _mask_ocr_text_lines(page: fitz.Page, lines: list[tuple[fitz.Rect, str]]) -> None:
    for rect, text in lines:
        if not text.strip():
            continue

        mask_rect = fitz.Rect(
            max(page.rect.x0, rect.x0 - PDF_OCR_TEXT_MASK_HORIZONTAL_PADDING),
            max(page.rect.y0, rect.y0 - PDF_OCR_TEXT_MASK_VERTICAL_PADDING),
            min(page.rect.x1, rect.x1 + PDF_OCR_TEXT_MASK_HORIZONTAL_PADDING),
            min(page.rect.y1, rect.y1 + PDF_OCR_TEXT_MASK_VERTICAL_PADDING),
        )
        page.draw_rect(mask_rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)


def _ocr_page_lines(page: fitz.Page, runtime: RuntimeSettings) -> list[tuple[fitz.Rect, str]]:
    pixmap = page.get_pixmap(dpi=PDF_OCR_DPI, alpha=False)
    image = _pixmap_to_image(pixmap)
    language = _ocr_language(runtime)
    if language:
        data = pytesseract.image_to_data(image, lang=language, output_type=pytesseract.Output.DICT)
    else:
        data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

    line_map: dict[tuple[int, int, int], dict[str, float | list[str]]] = {}
    scale_x = page.rect.width / pixmap.width
    scale_y = page.rect.height / pixmap.height

    for index, raw_text in enumerate(data["text"]):
        text = raw_text.strip()
        if not text:
            continue

        key = (data["block_num"][index], data["par_num"][index], data["line_num"][index])
        left = float(data["left"][index]) * scale_x
        top = float(data["top"][index]) * scale_y
        width = float(data["width"][index]) * scale_x
        height = float(data["height"][index]) * scale_y
        right = left + width
        bottom = top + height

        existing = line_map.get(key)
        if existing is None:
            line_map[key] = {"left": left, "top": top, "right": right, "bottom": bottom, "words": [text]}
            continue

        existing["left"] = min(existing["left"], left)
        existing["top"] = min(existing["top"], top)
        existing["right"] = max(existing["right"], right)
        existing["bottom"] = max(existing["bottom"], bottom)
        existing["words"].append(text)

    lines: list[tuple[fitz.Rect, str]] = []
    for key in sorted(line_map):
        line = line_map[key]
        rect = fitz.Rect(line["left"], line["top"], line["right"], line["bottom"])
        text = " ".join(line["words"])
        lines.append((rect, text))
    return lines


def _pdf_has_any_extractable_text(path: Path) -> bool:
    document = fitz.open(path)
    try:
        return any(_page_has_extractable_text(page) for page in document)
    finally:
        document.close()


def _page_luminance_metrics(page: fitz.Page) -> tuple[float, float]:
    pixmap = page.get_pixmap(dpi=PDF_BACKGROUND_ANALYSIS_DPI, alpha=False)
    samples = pixmap.samples
    pixel_count = len(samples) // 3
    if pixel_count == 0:
        return 0.0, 0.0

    total_luminance = 0.0
    bright_pixels = 0
    for index in range(0, len(samples), 3):
        red = samples[index]
        green = samples[index + 1]
        blue = samples[index + 2]
        luminance = (0.2126 * red) + (0.7152 * green) + (0.0722 * blue)
        total_luminance += luminance
        if luminance >= 245:
            bright_pixels += 1

    return total_luminance / pixel_count, bright_pixels / pixel_count


def _page_vector_content_metrics(page: fitz.Page) -> tuple[int, int]:
    drawing_items = 0
    line_drawings = 0
    for drawing in page.get_drawings():
        items = drawing.get("items", ())
        drawing_items += len(items)
        rect = drawing.get("rect")
        if rect and (rect.width < 2 or rect.height < 2):
            line_drawings += 1
    return drawing_items, line_drawings


def _pdf_prefers_ocr_workaround(path: Path) -> bool:
    document = fitz.open(path)
    try:
        sample_count = min(document.page_count, PDF_OCR_WORKAROUND_PAGE_SAMPLE_LIMIT)
        if sample_count == 0:
            return False

        total_luminance = 0.0
        total_bright_ratio = 0.0
        total_drawing_items = 0
        total_line_drawings = 0
        for page_index in range(sample_count):
            page = document[page_index]
            luminance, bright_ratio = _page_luminance_metrics(page)
            drawing_items, line_drawings = _page_vector_content_metrics(page)
            total_luminance += luminance
            total_bright_ratio += bright_ratio
            total_drawing_items += drawing_items
            total_line_drawings += line_drawings

        average_luminance = total_luminance / sample_count
        average_bright_ratio = total_bright_ratio / sample_count
        average_drawing_items = total_drawing_items / sample_count
        average_line_drawings = total_line_drawings / sample_count
        return (
            average_luminance >= PDF_OCR_WORKAROUND_MIN_LUMINANCE
            and average_bright_ratio >= PDF_OCR_WORKAROUND_MIN_BRIGHT_RATIO
            and average_drawing_items <= PDF_OCR_WORKAROUND_MAX_AVERAGE_DRAWING_ITEMS
            and average_line_drawings <= PDF_OCR_WORKAROUND_MAX_AVERAGE_LINE_DRAWINGS
        )
    finally:
        document.close()


def _prepare_pdf_for_babeldoc(
    input_path: str,
    prepared_path: Path,
    runtime: RuntimeSettings,
) -> tuple[int, bool]:
    document = fitz.open(input_path)
    try:
        has_text_by_page = [_page_has_extractable_text(page) for page in document]
        has_text_resources_by_page = [_page_has_text_resources(page) for page in document]
        page_count = document.page_count
        if all(has_text_by_page):
            return page_count, False

        # PDFs with broken or custom embedded fonts may not yield words via
        # PyMuPDF, but BabelDOC can still process their text objects correctly.
        # Treat them as native PDFs instead of rebuilding them through OCR.
        if any(has_text_resources_by_page):
            return page_count, False

        # Some native PDFs still contain a few pages that PyMuPDF cannot extract
        # words from reliably even though BabelDOC can translate the original
        # document well. Rebuilding the whole file to inject OCR text for those
        # pages degrades untouched native-text pages, especially dense tables.
        # Restrict OCR preprocessing to fully scanned/image-only PDFs.
        if any(has_text_by_page):
            return page_count, False

        if not runtime.ocr_enabled:
            raise RuntimeError("PDF contains scanned or image-only pages. Enable OCR to preserve layout during translation.")

        prepared_document = fitz.open()
        try:
            for index, has_text in enumerate(has_text_by_page):
                if has_text:
                    prepared_document.insert_pdf(document, from_page=index, to_page=index)
                    continue

                source_page = document[index]
                target_page = prepared_document.new_page(width=source_page.rect.width, height=source_page.rect.height)
                ocr_lines = _ocr_page_lines(source_page, runtime)
                target_page.show_pdf_page(source_page.rect, document, index)
                _mask_ocr_text_lines(target_page, ocr_lines)
                for rect, text in ocr_lines:
                    _insert_background_text_best_fit(target_page, rect, text)
            prepared_document.save(prepared_path)
        finally:
            prepared_document.close()

        if not _pdf_has_any_extractable_text(prepared_path):
            raise RuntimeError("OCR could not detect usable text in this PDF. Please try a clearer scan or a native-text PDF.")
        return page_count, True
    finally:
        document.close()


def _normalize_babeldoc_output_line(line: str) -> str:
    return " ".join(_ANSI_ESCAPE_RE.sub("", line).split())


def _map_babeldoc_overall_progress(overall_progress: float) -> int:
    ratio = min(max(overall_progress, 0.0), 100.0) / 100.0
    progress_span = _BABELDOC_PROGRESS_END - _BABELDOC_PROGRESS_START
    progress = _BABELDOC_PROGRESS_START + int(round(progress_span * ratio))
    return max(_BABELDOC_PROGRESS_START, min(_BABELDOC_PROGRESS_END, progress))


def _extract_babeldoc_progress_fraction(line: str) -> tuple[int, int] | None:
    matches = _BABELDOC_PROGRESS_FRACTION_RE.findall(line)
    if not matches:
        return None
    current, total = matches[-1]
    total_value = int(total)
    if total_value <= 0:
        return None
    return int(current), total_value


def _map_babeldoc_progress(stage: _BabeldocProgressStage, current: int, total: int) -> int:
    stage_ratio = min(max(current / total, 0.0), 1.0)
    overall_ratio = (stage.cumulative_weight_before + stage.weight * stage_ratio) / _BABELDOC_TOTAL_STAGE_WEIGHT
    progress_span = _BABELDOC_PROGRESS_END - _BABELDOC_PROGRESS_START
    progress = _BABELDOC_PROGRESS_START + int(round(progress_span * overall_ratio))
    return max(_BABELDOC_PROGRESS_START, min(_BABELDOC_PROGRESS_END, progress))


def _parse_babeldoc_progress_update(line: str) -> tuple[_BabeldocProgressStage, int, str] | None:
    normalized = _normalize_babeldoc_output_line(line)
    if not normalized:
        return None

    event_prefix_index = normalized.find(_BABELDOC_PROGRESS_EVENT_PREFIX)
    if event_prefix_index >= 0:
        payload = normalized[event_prefix_index + len(_BABELDOC_PROGRESS_EVENT_PREFIX) :]
        try:
            event = json.loads(payload)
        except json.JSONDecodeError:
            event = None
        if isinstance(event, dict):
            stage_name = event.get("stage")
            stage = _BABELDOC_PROGRESS_STAGE_BY_NAME.get(stage_name) if isinstance(stage_name, str) else None
            overall_progress = event.get("overall_progress")
            if stage is not None and isinstance(overall_progress, (int, float)):
                return stage, _map_babeldoc_overall_progress(float(overall_progress)), normalized

    for stage in _BABELDOC_PROGRESS_STAGES:
        if not normalized.startswith(stage.name):
            continue
        fraction = _extract_babeldoc_progress_fraction(normalized)
        progress = (
            _map_babeldoc_progress(stage, *fraction)
            if fraction is not None
            else _map_babeldoc_progress(stage, 0, 1)
        )
        return stage, progress, normalized
    return None


class _BabeldocProgressTracker:
    def __init__(self, session: Session, job: TranslationJob) -> None:
        self.session = session
        self.job = job
        self.last_status = job.status
        self.last_progress = job.progress
        self.last_stage_name: str | None = None

    def handle_output_line(self, line: str) -> None:
        parsed = _parse_babeldoc_progress_update(line)
        if parsed is None:
            return

        stage, progress, normalized_line = parsed
        progress = max(progress, self.last_progress)
        if (
            stage.status == self.last_status
            and progress == self.last_progress
            and stage.name == self.last_stage_name
        ):
            return

        stage_changed = stage.name != self.last_stage_name or stage.status != self.last_status
        update_job_state(
            self.session,
            self.job,
            status=stage.status,
            progress=progress,
            message=stage.message if stage_changed else None,
            details={"stage": stage.name, "line": normalized_line} if stage_changed else None,
        )
        self.last_status = stage.status
        self.last_progress = progress
        self.last_stage_name = stage.name


def _translate_prepared_pdf_with_babeldoc_hooks(
    input_path: Path,
    output_dir: Path,
    working_dir: Path,
    runtime: RuntimeSettings,
    job: TranslationJob,
    session: Session,
    *,
    source_language: str | None,
    target_language: str,
    use_ocr_workaround: bool,
) -> BabeldocLibraryResult:
    progress_tracker = _BabeldocProgressTracker(session, job)

    def on_progress_event(event: dict[str, object]) -> None:
        ensure_not_cancelled(session, job)
        progress_tracker.handle_output_line(
            _BABELDOC_PROGRESS_EVENT_PREFIX + json.dumps(event, ensure_ascii=False, separators=(",", ":"))
        )

    return translate_pdf_with_babeldoc_library(
        input_path,
        output_dir,
        working_dir,
        runtime,
        source_language=source_language,
        target_language=target_language,
        use_ocr_workaround=use_ocr_workaround,
        qps=BABELDOC_QPS,
        report_interval=BABELDOC_PROGRESS_REPORT_INTERVAL_SECONDS,
        on_progress_event=on_progress_event,
    )


def _paragraph_targets(document: Document) -> list:
    targets = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(paragraph for paragraph in cell.paragraphs if paragraph.text.strip())
    return targets


def replace_paragraph_text(paragraph, text: str) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            element.remove(child)
    paragraph.add_run(text)


def translate_docx(
    input_path: str,
    output_path: Path,
    translator: OpenAICompatibleTranslator,
    job: TranslationJob,
    session: Session,
) -> int | None:
    document = Document(input_path)
    targets = _paragraph_targets(document)
    texts = [paragraph.text for paragraph in targets]

    def on_progress(index: int, total: int) -> None:
        progress = 20 + int((index / max(total, 1)) * 60)
        update_job_state(session, job, status=JobStatus.TRANSLATING, progress=progress)

    translated = translate_segments(
        translator,
        texts,
        source_language=job.source_language,
        target_language=job.target_language,
        preserve_line_breaks=True,
        on_progress=on_progress,
        cancel_check=lambda: ensure_not_cancelled(session, job),
    )
    for paragraph, translated_text in zip(targets, translated, strict=True):
        replace_paragraph_text(paragraph, translated_text)
    document.save(output_path)
    return None


def _restore_input_file_from_duplicate(session: Session, input_file: JobFile, target_path: Path) -> Path | None:
    candidates = (
        session.query(JobFile)
        .filter(JobFile.kind == JobFileKind.INPUT)
        .filter(JobFile.checksum == input_file.checksum)
        .filter(JobFile.id != input_file.id)
        .filter(JobFile.deleted_at.is_(None))
        .order_by(JobFile.created_at.desc())
        .all()
    )
    for candidate in candidates:
        candidate_path = Path(candidate.storage_path)
        try:
            if candidate_path.stat().st_size != input_file.size_bytes:
                continue
        except FileNotFoundError:
            continue
        if file_checksum(candidate_path) != input_file.checksum:
            continue

        target_path.parent.mkdir(parents=True, exist_ok=True)
        restore_path = target_path.with_name(f".{target_path.name}.restore")
        try:
            shutil.copyfile(candidate_path, restore_path)
            if restore_path.stat().st_size != input_file.size_bytes or file_checksum(restore_path) != input_file.checksum:
                restore_path.unlink(missing_ok=True)
                continue
            restore_path.replace(target_path)
        finally:
            restore_path.unlink(missing_ok=True)
        return candidate_path
    return None


def _ensure_input_file_available(session: Session, job: TranslationJob) -> Path:
    input_file = job.input_file
    input_path = Path(input_file.storage_path)
    try:
        actual_size = input_path.stat().st_size
    except FileNotFoundError:
        actual_size = None

    if actual_size == input_file.size_bytes and actual_size > 0:
        return input_path

    restored_from = _restore_input_file_from_duplicate(session, input_file, input_path)
    if restored_from is not None:
        add_job_event(
            session,
            job,
            "Restored missing or incomplete input file from matching upload",
            level="warning",
            details={"restored_from": str(restored_from), "expected_bytes": input_file.size_bytes, "actual_bytes": actual_size},
        )
        session.commit()
        return input_path

    actual_text = "missing" if actual_size is None else f"{actual_size} bytes"
    raise RuntimeError(
        f"Input file is unavailable or incomplete: expected {input_file.size_bytes} bytes at {input_path}, found {actual_text}. "
        "Please re-upload the source document."
    )


def translate_pdf(
    input_path: str,
    output_path: Path,
    runtime: RuntimeSettings,
    job: TranslationJob,
    session: Session,
) -> int:
    source_language = _babeldoc_language_code(job.source_language, allow_auto=True)
    target_language = _babeldoc_language_code(job.target_language, allow_auto=False)

    with tempfile.TemporaryDirectory(prefix="babeldoc-") as temp_dir:
        temp_root = Path(temp_dir)
        source_input_path = temp_root / f"source{Path(input_path).suffix.lower()}"
        prepared_input_path = temp_root / f"prepared{Path(input_path).suffix.lower()}"
        output_dir = temp_root / "output"
        working_dir = temp_root / "working"
        output_dir.mkdir(parents=True, exist_ok=True)
        working_dir.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(input_path, source_input_path)

        page_count, used_ocr = _prepare_pdf_for_babeldoc(str(source_input_path), prepared_input_path, runtime)
        babeldoc_input_path = prepared_input_path if used_ocr else source_input_path
        # OCR workaround paints over original content before re-typesetting it.
        # Use it for OCR-prepared PDFs and white-background native PDFs where
        # BabelDOC may otherwise leave the original text visible under the translation.
        use_ocr_workaround = used_ocr or _pdf_prefers_ocr_workaround(babeldoc_input_path)

        if used_ocr:
            update_job_state(session, job, status=JobStatus.OCR_RUNNING, progress=18, message="Prepared searchable PDF for scanned pages")

        update_job_state(
            session,
            job,
            status=JobStatus.PARSING,
            progress=20,
            message="Starting layout-preserving PDF translation",
            details={
                "ocr_workaround": use_ocr_workaround,
                "ocr_prepared": used_ocr,
                "pdf_mode": "babeldoc_library_hooks",
                "enhance_compatibility": False,
            },
        )
        babeldoc_result = _translate_prepared_pdf_with_babeldoc_hooks(
            babeldoc_input_path,
            output_dir,
            working_dir,
            runtime,
            job,
            session,
            source_language=source_language,
            target_language=target_language,
            use_ocr_workaround=use_ocr_workaround,
        )

        update_job_state(session, job, status=JobStatus.REBUILDING, progress=92, message="Saving translated PDF")
        shutil.move(str(babeldoc_result.mono_output), output_path)
        if babeldoc_result.hook_sidecar and babeldoc_result.hook_sidecar.exists():
            sidecar_output_path = babeldoc_ir_sidecar_path(output_path)
            shutil.copyfile(babeldoc_result.hook_sidecar, sidecar_output_path)
            add_job_event(
                session,
                job,
                "Recorded BabelDOC internal hook IR",
                details={"sidecar": str(sidecar_output_path)},
            )

    output_document = fitz.open(output_path)
    try:
        return output_document.page_count
    finally:
        output_document.close()


def test_model_connection(runtime: RuntimeSettings) -> tuple[int, str]:
    translator = OpenAICompatibleTranslator(runtime)
    return translator.test_connection()


def run_translation_job(job_id: str) -> None:
    session = SessionLocal()
    try:
        job = (
            session.query(TranslationJob)
            .options(
                selectinload(TranslationJob.input_file),
                selectinload(TranslationJob.output_file),
                selectinload(TranslationJob.created_by_user),
            )
            .filter(TranslationJob.id == job_id)
            .first()
        )
        if job is None:
            logger.warning("Job not found", extra={"job_id": job_id})
            return
        if job.status == JobStatus.CANCELLED:
            return

        runtime = get_runtime_settings(session)
        update_job_state(session, job, status=JobStatus.PARSING, progress=5, message="Parsing document")
        ensure_not_cancelled(session, job)

        input_file = job.input_file
        input_path = str(_ensure_input_file_available(session, job))
        extension = Path(input_file.original_name).suffix.lower()
        output_path = build_output_target(runtime, input_file.original_name, extension)
        page_count: int | None

        if extension == ".pdf":
            page_count = translate_pdf(input_path, output_path, runtime, job, session)
        elif extension == ".docx":
            translator = OpenAICompatibleTranslator(runtime)
            update_job_state(session, job, status=JobStatus.TRANSLATING, progress=20, message="Translating DOCX content")
            page_count = translate_docx(input_path, output_path, translator, job, session)
            update_job_state(session, job, status=JobStatus.REBUILDING, progress=88, message="Writing translated DOCX")
        else:
            raise RuntimeError("Unsupported file type")

        output_file = JobFile(
            original_name=output_path.name,
            stored_name=output_path.name,
            storage_path=str(output_path),
            content_type=input_file.content_type,
            size_bytes=output_path.stat().st_size,
            checksum=file_checksum(output_path),
            kind=JobFileKind.OUTPUT,
            created_by=job.created_by,
        )
        session.add(output_file)
        session.flush()

        job.output_file_id = output_file.id
        job.output_file = output_file
        job.page_count = page_count
        preview_details: dict | None = None
        try:
            preview = load_or_create_preview(job, force=True)
            preview_details = {"preview_pages": len(preview["pages"])}
        except Exception as exc:
            if extension == ".pdf":
                raise
            logger.warning("Preview preparation failed", extra={"job_id": job.id, "error": str(exc)})
            add_job_event(session, job, "Preview could not be prepared", details={"error": str(exc)})
        update_job_state(session, job, status=JobStatus.COMPLETED, progress=100, message="Translation completed")
        record_audit(
            session,
            action="jobs.completed",
            entity_type="translation_job",
            entity_id=job.id,
            actor_id=job.created_by,
            details={"status": job.status.value, "output_file_id": output_file.id, **(preview_details or {})},
        )
        session.commit()
        logger.info("Completed translation job", extra={"job_id": job.id})
    except JobCancelledError:
        if "job" in locals():
            update_job_state(session, job, status=JobStatus.CANCELLED, progress=job.progress, message="Job cancelled")
            record_audit(
                session,
                action="jobs.cancelled",
                entity_type="translation_job",
                entity_id=job.id,
                actor_id=job.created_by,
                details={"status": job.status.value},
            )
            session.commit()
    except Exception as exc:
        if "job" in locals():
            logger.exception("Translation job failed", extra={"job_id": job.id})
            update_job_state(
                session,
                job,
                status=JobStatus.FAILED,
                progress=job.progress,
                error_message=str(exc),
                message="Translation failed",
                details={"error": str(exc)},
            )
            record_audit(
                session,
                action="jobs.failed",
                entity_type="translation_job",
                entity_id=job.id,
                actor_id=job.created_by,
                details={"error": str(exc)},
            )
            session.commit()
        else:
            logger.exception("Translation job failed before loading state", extra={"job_id": job_id})
    finally:
        session.close()
